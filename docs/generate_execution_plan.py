"""生成企业内部 Agent 智能体平台执行计划 Word 文档。

Usage:
    python generate_execution_plan.py

Output:
    /mnt/e/program/agent_platform/output/Agent平台执行计划.docx
"""
import sys
from pathlib import Path

# 复用 generate_docx.py 里的样式函数
sys.path.insert(0, str(Path(__file__).parent))
from generate_docx import (
    set_cell_bg, set_run_font, set_paragraph_spacing,
    add_page_break, add_horizontal_line, set_table_borders,
    init_document, add_h1, add_h2, add_h3, add_h4,
    add_body, add_bullet, add_number, add_code_block, add_quote,
    add_data_table,
)

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================================
# 封面（执行计划版）
# ============================================================================

def add_cover(doc):
    """执行计划封面。"""
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=18, line_spacing=1.5)
    run = p.add_run("企业内部 Agent 智能体平台")
    set_run_font(run, name_cn="黑体", size_pt=30, bold=True, color="1F4E79")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)
    run = p.add_run("项 目 执 行 计 划")
    set_run_font(run, name_cn="黑体", size_pt=24, bold=True, color="2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=36, line_spacing=1.5)
    run = p.add_run("Project Execution Plan")
    set_run_font(run, size_pt=14, color="808080")

    # 分隔线
    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sep_p, before=0, after=24, line_spacing=1.0)
    pPr = sep_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=24, after=12, line_spacing=1.5)
    run = p.add_run("—— MVP 阶段（10 周）实施与交付计划 ——")
    set_run_font(run, size_pt=14, color="404040")

    for _ in range(2):
        doc.add_paragraph()

    # 元信息表
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT_CENTER if False else 1
    from docx.enum.table import WD_TABLE_ALIGNMENT
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("项目名称", "企业内部 Agent 智能体平台"),
        ("项目类型", "企业内部 IT 基础平台建设"),
        ("项目周期", "10 周（MVP 阶段）"),
        ("文档版本", "V1.0"),
        ("编制日期", "2026-08-28"),
        ("文档密级", "内部公开"),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c1 = row.cells[0]
        c2 = row.cells[1]
        c1.width = Cm(4)
        c2.width = Cm(9)
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(k)
        set_run_font(run1, name_cn="黑体", size_pt=12, bold=True, color="FFFFFF")
        from docx.enum.table import WD_ALIGN_VERTICAL
        set_cell_bg(c1, "1F4E79")
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(v)
        set_run_font(run2, size_pt=12, color="1F4E79")
        set_cell_bg(c2, "F2F7FB")
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(meta_table, color="1F4E79", size="8")

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=4, line_spacing=1.5)
    run = p.add_run("编制单位：企业数字化转型办公室")
    set_run_font(run, size_pt=13, bold=True, color="404040")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)
    run = p.add_run("2026 年 8 月")
    set_run_font(run, size_pt=12, color="666666")

    add_page_break(doc)


# ============================================================================
# 目录
# ============================================================================

def add_toc(doc):
    """目录页。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=18, line_spacing=1.5)
    run = p.add_run("目  录")
    set_run_font(run, name_cn="黑体", size_pt=20, bold=True, color="1F4E79")
    add_horizontal_line(p)

    toc_entries = [
        ("第一章  概述", "1"),
        ("  1.1 文档目的", "1"),
        ("  1.2 适用范围", "1"),
        ("  1.3 与其他文档的关系", "1"),
        ("  1.4 执行原则", "2"),
        ("第二章  项目概况", "3"),
        ("  2.1 项目背景", "3"),
        ("  2.2 项目目标", "3"),
        ("  2.3 关键决策摘要", "4"),
        ("  2.4 范围界定", "4"),
        ("第三章  团队与角色", "6"),
        ("  3.1 团队组成", "6"),
        ("  3.2 角色与职责", "6"),
        ("  3.3 沟通机制", "7"),
        ("第四章  阶段划分与任务分解", "8"),
        ("  4.1 总体时间线", "8"),
        ("  4.2 阶段 0：仓库骨架与基础设施（1 周）", "9"),
        ("  4.3 阶段 1：后端鉴权与权限模型（1.5 周）", "10"),
        ("  4.4 阶段 2：Dify 适配 + 同步 + 授权（1 周）", "12"),
        ("  4.5 阶段 3：对话代理 + SSE + 对话镜像（2 周）", "13"),
        ("  4.6 阶段 4：文件上传（0.5 周）", "15"),
        ("  4.7 阶段 5：前端骨架 + 登录 + 对话 UI（2 周）", "16"),
        ("  4.8 阶段 6：管理员后台（1 周）", "17"),
        ("  4.9 阶段 7：生产部署 + 可观测性（1 周）", "18"),
        ("第五章  任务依赖与关键路径", "20"),
        ("  5.1 任务依赖关系", "20"),
        ("  5.2 关键路径分析", "20"),
        ("  5.3 并行执行机会", "21"),
        ("第六章  里程碑与交付物", "22"),
        ("  6.1 里程碑表", "22"),
        ("  6.2 交付物清单", "23"),
        ("第七章  资源与预算", "25"),
        ("  7.1 人力资源", "25"),
        ("  7.2 硬件资源", "25"),
        ("  7.3 第三方服务与软件", "26"),
        ("第八章  风险与应对", "27"),
        ("  8.1 风险登记册", "27"),
        ("  8.2 风险监控与升级", "29"),
        ("第九章  质量保证", "30"),
        ("  9.1 质量目标", "30"),
        ("  9.2 评审机制", "30"),
        ("  9.3 测试策略", "31"),
        ("第十章  变更管理", "32"),
        ("  10.1 变更流程", "32"),
        ("  10.2 变更控制委员会", "32"),
        ("第十一章  沟通与汇报", "33"),
        ("  11.1 会议节奏", "33"),
        ("  11.2 报告机制", "33"),
        ("  11.3 升级机制", "34"),
        ("附录 A  阶段 0–4 任务详细清单", "35"),
        ("附录 B  关键路径与甘特图说明", "38"),
        ("附录 C  文档索引", "39"),
    ]
    for title, page in toc_entries:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line_spacing=1.4)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
        run = p.add_run(title)
        if title.startswith("第") and "章" in title[:3]:
            set_run_font(run, name_cn="黑体", size_pt=11, bold=True)
        else:
            set_run_font(run, size_pt=10.5)
        tab_run = p.add_run("\t")
        page_run = p.add_run(page)
        set_run_font(page_run, size_pt=10.5)

    add_page_break(doc)


# ============================================================================
# 各章节内容
# ============================================================================

def write_chapter_1(doc):
    add_h1(doc, "第一章  概述")

    add_h2(doc, "1.1 文档目的")
    add_body(doc,
        "本文档定义企业内部 Agent 智能体平台 MVP 阶段的执行计划，"
        "明确项目目标、范围、阶段划分、任务分解、团队角色、里程碑、风险及质量保证等关键要素，"
        "为项目团队提供统一的工作指引，并作为项目干系人对齐预期、跟踪进度的基准文档。")

    add_body(doc, "本文档具体用于：")
    add_bullet(doc, "项目团队按阶段、按任务有序开展工作；")
    add_bullet(doc, "项目经理跟踪进度、识别风险、汇报状态；")
    add_bullet(doc, "项目干系人（业务方、IT 负责人、安全 review 等）评审与决策；")
    add_bullet(doc, "变更控制委员会（CCB）评估范围、进度、成本变更的基线。")

    add_h2(doc, "1.2 适用范围")
    add_body(doc, "本计划适用于 MVP 阶段（10 周）的实施工作，覆盖：")
    add_bullet(doc, "后端 FastAPI 服务（鉴权、权限、Dify 集成、对话代理、文件、审计、加密导出）；")
    add_bullet(doc, "前端 React SPA（登录、对话、历史、管理员后台、解密页面）；")
    add_bullet(doc, "基础设施（Postgres、Dify 社区版、Nginx、Docker Compose）；")
    add_bullet(doc, "测试与质量保证（单测、集成、E2E、手动验证）；")
    add_bullet(doc, "生产部署与可观测性（仅 MVP 范围内的最小集）。")

    add_body(doc, "以下内容不在本文档范围内：")
    add_bullet(doc, "二期 / 三期特性（限流、配额、SSO、多模态等）——待后续独立 plan 文件描述；")
    add_bullet(doc, "业务 Agent 的具体设计与编排——由 Dify 端业务团队负责；")
    add_bullet(doc, "硬件采购流程——由 IT 采购部门按公司流程执行。")

    add_h2(doc, "1.3 与其他文档的关系")
    add_data_table(doc,
        headers=["文档类型", "文档名", "作用"],
        rows=[
            ["设计文档",
             "企业内部 Agent 智能体平台设计文档",
             "本执行计划的输入：定义做什么"],
            ["环境准备",
             "实施环境准备文档",
             "阶段 0 的具体操作步骤：怎么搭环境"],
            ["技术方案",
             "建设技术方案框架",
             "本执行计划的总体设计依据：技术选型 + 关键设计"],
            ["加密导出",
             "加密导出特性设计文档",
             "阶段 3 后半段 / 阶段 4 的输入"],
            ["执行计划（本文件）",
             "项目执行计划",
             "团队工作指南：何时做、谁来做、交付什么"],
            ["运维手册（二期）",
             "上线运行手册",
             "上线后由 SRE 维护，不在 MVP 范围"],
        ],
        col_widths=[3.0, 5.5, 6.8],
        first_col_bold=True,
    )

    add_quote(doc,
        "本执行计划是\"何时做、谁来做\"的指南，技术方案框架是\"做什么、怎么做\"的依据。\n"
        "实施过程中如对设计产生疑问，先查技术方案框架；如对节奏有疑问，先查本计划。")

    add_h2(doc, "1.4 执行原则")
    add_data_table(doc,
        headers=["原则", "含义"],
        rows=[
            ["小步快跑，持续集成", "每个 Task 都跑通测试再进入下一个；每周末整合可演示版本"],
            ["TDD 优先", "关键模块（鉴权、授权、SSE 透传）先写测试再写实现"],
            ["配置即代码", "所有环境变量、密钥、Docker 配置进 git；密钥不进 git"],
            ["安全左移", "鉴权、CSRF、文件上传、加密导出的安全要求在对应阶段第一时间实现"],
            ["可观测先行", "关键路径日志结构化 + 关键审计动作落库；不推迟到上线前"],
            ["文档同步", "代码与文档同 PR；变更影响文档时必须同步更新"],
            ["范围刚性", "MVP 范围外特性不接受口头加塞；走变更管理流程"],
            ["风险显性化", "风险每日 stand-up 同步；达到阈值立即升级"],
        ],
        col_widths=[4.0, 11.3],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_chapter_2(doc):
    add_h1(doc, "第二章  项目概况")

    add_h2(doc, "2.1 项目背景")
    add_body(doc,
        "随着大语言模型在企业内部多场景的快速渗透，"
        "公司在文档检索、IT 工单、HR 答疑、技术问答等高频场景中员工日均产生大量重复性问题，"
        "严重占用业务专家时间。现有方案存在三大痛点：分散的 AI 入口、缺乏权限隔离、审计与合规困难。")

    add_body(doc,
        "本项目旨在构建一个企业内部的 Agent 智能体平台，统一对外提供 AI 对话入口，"
        "并通过自研的鉴权 + 代理 + 审计层，将底层 LLM 引擎（Dify 社区版）的能力"
        "以安全、可控、可审计的方式开放给公司全体员工。")

    add_h2(doc, "2.2 项目目标")

    add_h3(doc, "2.2.1 总体目标")
    add_data_table(doc,
        headers=["目标维度", "具体目标", "衡量指标"],
        rows=[
            ["统一入口", "面向员工提供统一的 Web 门户", "Portal 上线，员工登录可达"],
            ["多 Agent 复用", "支持对接 5+ 业务 Agent（HR、IT、研发知识库等）", "管理员可自助授权"],
            ["权限隔离", "按用户 / 部门 / 角色三态授权", "权限校验覆盖率达 100%"],
            ["对话体验", "首字节延迟 ≤ 300ms", "P95 首字节延迟 ≤ 500ms"],
            ["审计合规", "所有敏感操作有日志，可导出", "审计覆盖关键路径 100%"],
            ["文件安全", "支持常见办公文档 / 图片上传", "大小 / MIME / 文件名三重校验"],
            ["数据自主", "会话可加密导出 / 解密", "密码不出客户端"],
        ],
        col_widths=[2.8, 7.0, 4.5],
        first_col_bold=True,
    )

    add_h3(doc, "2.2.2 MVP 范围")
    add_body(doc, "MVP 阶段（10 周）的核心交付：")
    add_bullet(doc, "Dify 部署调通（外部化 Postgres / Redis、镜像版本锁定）")
    add_bullet(doc, "FastAPI 骨架 + JWT 登录 + CSRF 防护 + 密钥拆分")
    add_bullet(doc, "对话流（含 SSE 透传 + 错误兜底 + 写库时机）")
    add_bullet(doc, "多 Agent 授权（用户 / 部门 / 角色三态）")
    add_bullet(doc, "对话镜像 + 审计日志")
    add_bullet(doc, "文件上传（大小 / MIME / 文件名安全约束）")
    add_bullet(doc, "基础管理员后台：用户管理 + Agent 授权")
    add_bullet(doc, "加密导出（浏览器侧 AES-256-GCM + PBKDF2 + /decrypt 页面）")
    add_bullet(doc, "生产部署最小集（Docker Compose + Nginx + 日志）")

    add_h2(doc, "2.3 关键决策摘要")
    add_data_table(doc,
        headers=["决策项", "决策", "影响"],
        rows=[
            ["底层引擎", "复用 Dify 社区版", "放弃自研 Agent 引擎；运维需熟悉 Dify"],
            ["后端框架", "FastAPI（Python 3.11+）", "团队 Python 能力匹配；异步生态完善"],
            ["前端框架", "React 18 + TypeScript + Vite + AntD 5", "公司技术栈匹配；招聘相对容易"],
            ["数据库", "PostgreSQL 16", "JSONB / 物化路径 / 异步适配良好"],
            ["鉴权方案", "JWT 双 token + httpOnly + SameSite=Strict", "无状态；refresh 主动撤销"],
            ["部署形态", "Docker Compose 一把梭", "环境一致；运维成本低"],
            ["加密导出", "浏览器侧 Web Crypto + AES-256-GCM", "密码不出客户端；支持 /decrypt 离线"],
            ["MVP 不引 Redis", "JWT 无状态；限流二期", "避免单点；二期需重新设计限流架构"],
        ],
        col_widths=[3.0, 5.5, 6.8],
        first_col_bold=True,
    )

    add_h2(doc, "2.4 范围界定")

    add_h3(doc, "2.4.1 MVP 包含")
    add_data_table(doc,
        headers=["模块", "范围"],
        rows=[
            ["用户与鉴权", "邮箱 + 密码登录、JWT 双 token、CSRF 防护、refresh 轮转"],
            ["组织架构", "用户 / 部门 / 角色基础 CRUD + 物化路径"],
            ["Agent 同步", "启动时从 Dify 全量拉取应用 + 加密存储 API Key"],
            ["Agent 授权", "用户 / 部门 / 角色三态授权 + 权限解析函数"],
            ["对话代理", "SSE 流式透传 + 错误兜底 + finally 写库 + 限流（最小集）"],
            ["对话镜像", "UUID 主键 + token_usage JSONB + 软删"],
            ["文件上传", "20MB + MIME 白名单 + 文件名清洗 + 审计"],
            ["加密导出", "单 / 批量会话 + 审计日志 + 附件，浏览器侧加密 + /decrypt 页面"],
            ["前端", "登录 / 对话 / 历史 / 管理员后台（用户管理 + Agent 授权）"],
            ["生产部署", "Docker Compose + Nginx + 镜像版本锁定 + 健康检查"],
            ["可观测性", "结构化日志（loguru）+ 关键审计动作落库；ELK / Prometheus 二期"],
        ],
        col_widths=[3.0, 12.3],
        first_col_bold=True,
    )

    add_h3(doc, "2.4.2 MVP 不包含")
    add_data_table(doc,
        headers=["模块", "推迟到"],
        rows=[
            ["限流（Redis 令牌桶）", "二期"],
            ["用量配额（每用户 / 部门每月 token 上限）", "二期"],
            ["用户反馈（👍/👎）", "二期"],
            ["管理员统计视图", "二期"],
            ["SSE 断点续传", "二期"],
            ["离职联动禁用 Dify 账号的自动化脚本", "二期（接口先预留）"],
            ["pgvector 跨对话检索", "二期"],
            ["部门 / 角色完整 CRUD UI", "MVP 基础版已包含，二期补全功能"],
            ["ClamAV 病毒扫描", "二期"],
            ["Sensitive Word 扫描", "二期"],
            ["企业 SSO / LDAP", "三期"],
            ["多模态上传预览", "三期"],
            ["团队协作（共享对话）", "三期"],
            ["PDF / Word 格式导出", "二期"],
        ],
        col_widths=[6.0, 9.3],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_chapter_3(doc):
    add_h1(doc, "第三章  团队与角色")

    add_h2(doc, "3.1 团队组成")
    add_body(doc, "MVP 阶段核心团队编制（4 人 + 兼职）：")
    add_data_table(doc,
        headers=["角色", "人数", "来源", "投入度"],
        rows=[
            ["项目经理（PM）", "1", "数字化办公室", "全程（10 周）"],
            ["后端工程师 BE-1（主程）", "1", "应用开发组", "全程（10 周）"],
            ["后端工程师 BE-2（Dify / SSE）", "1", "应用开发组", "全程（10 周）"],
            ["前端工程师 FE-1", "1", "前端开发组", "阶段 5-7 全程 + 阶段 0-4 兼职支持"],
            ["SRE / DevOps", "0.3（兼）", "基础架构组", "阶段 0 + 阶段 7 集中投入"],
            ["安全 review", "0.2（兼）", "安全团队", "阶段 1 / 3 / 4 关键节点"],
            ["测试工程师（兼）", "0.5（兼）", "QA 组", "阶段 3 末 + 阶段 6 / 7 集中投入"],
            ["业务代表（产品）", "0.3（兼）", "数字化办公室", "全程"],
        ],
        col_widths=[4.5, 2.0, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h2(doc, "3.2 角色与职责")
    add_data_table(doc,
        headers=["角色", "主要职责", "决策权"],
        rows=[
            ["项目经理", "进度跟踪、风险升级、跨组协调、汇报材料",
             "进度 / 范围微调；重大变更需 CCB"],
            ["后端主程", "架构设计、核心代码 review、鉴权 / 权限模块、安全 review 配合",
             "技术方案微调；接口契约"],
            ["后端工程师（Dify / SSE）", "Dify 集成、SSE 透传、文件上传、加密导出后端实现",
             "实现细节；性能优化"],
            ["前端工程师", "React SPA、对话 UI、管理后台、加密导出前端、/decrypt 页面",
             "前端架构；UI 交互细节"],
            ["SRE / DevOps", "Docker 编排、Nginx 配置、生产部署、监控告警",
             "基础设施选型；部署流程"],
            ["安全 review", "鉴权 / CSRF / 文件上传 / 加密导出的安全评审",
             "安全策略；威胁建模"],
            ["测试工程师", "测试用例设计、E2E 自动化、手动验证清单执行",
             "测试方法；缺陷验证"],
            ["业务代表", "需求澄清、UAT 验收、跨部门协调",
             "业务需求优先级；验收标准"],
        ],
        col_widths=[3.5, 9.0, 2.8],
        first_col_bold=True,
    )

    add_h2(doc, "3.3 沟通机制")
    add_data_table(doc,
        headers=["会议", "频率", "时长", "参与者", "目的"],
        rows=[
            ["每日 stand-up", "每日 09:30", "15 分钟", "全体开发", "昨日进展 / 今日计划 / 阻塞"],
            ["周计划会", "周一 10:00", "1 小时", "PM + 全体开发", "周计划 / 风险同步"],
            ["周评审会", "周五 16:00", "1 小时", "全体 + 业务代表", "本周交付演示 / 验收"],
            ["阶段评审", "每个阶段末", "2 小时", "全体 + 干系人", "阶段成果评审 / 阶段准入"],
            ["月度汇报", "每月最后一周五", "1 小时", "PM + 干系人", "月度状态汇报 / 决策项"],
            ["临时专题会", "按需", "30-60 分钟", "相关人员", "专项问题（如安全评审）"],
        ],
        col_widths=[3.0, 2.5, 1.8, 3.5, 4.5],
        first_col_bold=True,
    )

    add_quote(doc,
        "沟通原则：\n"
        "- 阻塞立即升级，不堆积；\n"
        "- 会议必须有结论 + 行动项 + 责任人；\n"
        "- 远程为主，本地办公区每周 1 次集中（周五评审日）；\n"
        "- IM 工作时间内 30 分钟内响应。")

    add_page_break(doc)


def write_chapter_4(doc):
    add_h1(doc, "第四章  阶段划分与任务分解")

    add_h2(doc, "4.1 总体时间线")
    add_body(doc, "本项目 MVP 阶段预计 10 周，分 8 个阶段（0–7）交付：")
    add_data_table(doc,
        headers=["阶段", "周次", "时长", "主要交付", "关键里程碑"],
        rows=[
            ["阶段 0", "W1", "1 周", "仓库骨架 + Docker + Dify 部署", "M1 环境就绪"],
            ["阶段 1", "W2 - W3 中", "1.5 周", "鉴权 + 权限模型", "M2 鉴权可用"],
            ["阶段 2", "W3 中 - W4 中", "1 周", "Dify 适配 + Agent 同步 + 授权", "M3 授权生效"],
            ["阶段 3", "W4 中 - W6 中", "2 周", "对话代理 + SSE + 对话镜像 + 加密导出后端", "M4 对话可用"],
            ["阶段 4", "W6 中 - W6 末", "0.5 周", "文件上传", "M5 文件上传"],
            ["阶段 5", "W7 - W8", "2 周", "前端骨架 + 登录 + 对话 UI + 加密导出前端", "M6 前端 MVP"],
            ["阶段 6", "W9", "1 周", "管理员后台", "M7 管理后台"],
            ["阶段 7", "W10", "1 周", "生产部署 + 可观测性", "M8 上线"],
            ["合计", "10 周", "10 周", "—", "—"],
        ],
        col_widths=[1.5, 2.0, 1.5, 6.5, 3.8],
        first_col_bold=True,
    )

    add_quote(doc,
        "原计划 4–6 周过于乐观——Dify 生产部署、外部化 PG / Redis、SSE 三层联调、"
        "多态授权一致性、CSRF / 文件上传安全均需工时。实际估时 10 周。")

    add_h2(doc, "4.2 阶段 0：仓库骨架与基础设施（1 周）")
    add_body(doc, "目标：搭建可运行的仓库骨架，部署 Dify 社区版 + Postgres，能用 curl 调通 Dify 管理 API。")

    add_h3(doc, "Task 0.1：初始化仓库骨架")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["git init + .gitignore", "BE-1", "git 仓库可提交"],
            ["backend/pyproject.toml + 依赖清单", "BE-1", "可 pip install -e .[dev]"],
            ["backend/.env.example", "BE-1", "可 cp .env.example .env"],
            ["backend/app/main.py 最小骨架", "BE-1", "/api/health 返回 200"],
            ["frontend/package.json + vite.config + tsconfig", "FE-1", "npm run dev 起前端"],
            ["deploy/docker-compose.yml + nginx/conf", "SRE", "docker compose up 起来"],
            ["deploy/Makefile", "SRE", "make up / down 一键操作"],
            ["README.md", "PM", "团队上手文档"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 0.2：部署 Dify 社区版 + 验证管理 API")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["克隆 Dify 1.1.0 + 外部化 Postgres", "SRE", "Dify 跑在外部 PG"],
            ["Dify 容器全部 healthy", "SRE", "docker compose ps 全绿"],
            ["初始化 Dify 管理员（/install）", "SRE", "可登录 Dify Web"],
            ["验证 /v1/setup/initialize-status", "SRE", "返回 finished: true"],
            ["Dify 中创建测试 Agent（test-agent）", "SRE", "拿到 API Key"],
            ["验证应用 API 调通（curl 流式）", "SRE", "看到 SSE event:message"],
            ["DIFY_DEPLOY.md 记录版本 + 凭据位置", "SRE", "凭据入 1Password"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 0 准入标准")
    add_bullet(doc, "git 仓库可提交；")
    add_bullet(doc, "Postgres 容器 + Dify 容器均 healthy；")
    add_bullet(doc, "curl http://localhost/v1/setup/initialize-status 返回 finished:true；")
    add_bullet(doc, "curl Dify /v1/chat-messages 看到 SSE 流；")
    add_bullet(doc, "BE-1 / FE-1 / SRE 各自能跑通开发环境。")

    add_h2(doc, "4.3 阶段 1：后端鉴权与权限模型（1.5 周）")
    add_body(doc, "目标：用户可注册/登录，颁发 JWT，CSRF 防护到位，权限模型数据表就位。")

    add_h3(doc, "Task 1.1：配置 + 数据库连接 + Alembic 迁移")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/core/config.py（pydantic-settings）", "BE-1", "可 get_settings() 读 .env"],
            ["app/db/{base,session,deps}.py", "BE-1", "async engine + get_db 依赖"],
            ["app/models/{user,department,role,refresh_token}.py", "BE-1", "5 张表 ORM"],
            ["alembic init + 首个迁移", "BE-1", "alembic upgrade head 成功"],
            ["测试库 + conftest.py（session + client fixture）", "BE-1", "pytest 通过"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 1.2：JWT 签发与校验")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/core/security.py（hash + JWT + refresh hash）", "BE-1", "3 个单测通过"],
            ["密码 bcrypt 哈希", "BE-1", "verify_password 正确"],
            ["access token 15min + refresh token 7d", "BE-1", "载荷含 jti / exp"],
            ["refresh token SHA256 哈希入库", "BE-1", "原始 token 不入库"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 1.3：认证 API + CSRF 中间件")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/schemas/auth.py", "BE-1", "LoginRequest / MeResponse"],
            ["app/auth/service.py（authenticate / issue_tokens / rotate）", "BE-1", "refresh 轮转生效"],
            ["app/auth/router.py（login / logout / refresh / me）", "BE-1", "cookie httponly + samesite=strict"],
            ["app/auth/deps.py（current_user）", "BE-1", "Depends(current_user) 可用"],
            ["app/core/middleware.py（CSRFMiddleware）", "BE-1", "跨域写请求 403"],
            ["test_auth.py 全套", "BE-1", "6 个测试用例通过"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 1.4：组织架构 CRUD（部门 / 角色）")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/models/{department,role}.py 完善", "BE-1", "物化路径 + CheckConstraint"],
            ["app/schemas/{dept,role}.py", "BE-1", "Create / Out DTO"],
            ["app/depts/router.py + app/roles/router.py", "BE-1", "POST /api/admin/departments 等"],
            ["app/admin/router.py 聚合", "BE-1", "PLATFORM_ADMIN 守卫"],
            ["test_admin.py", "BE-1", "PLATFORM_ADMIN CRUD 通过"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 1 准入标准")
    add_bullet(doc, "登录 → 拿到 httpOnly cookie；")
    add_bullet(doc, "GET /api/auth/me 携带 cookie 返回用户信息；")
    add_bullet(doc, "跨域 POST 被 CSRFMiddleware 拒绝（403）；")
    add_bullet(doc, "refresh 轮转：旧 refresh 二次使用 401；")
    add_bullet(doc, "PLATFORM_ADMIN 可创建部门 / 角色；")
    add_bullet(doc, "pytest 全套通过。")

    add_h2(doc, "4.4 阶段 2：Dify 适配 + 同步 + 授权（1 周）")
    add_body(doc, "目标：Dify 应用镜像到本地 DB，按用户/部门/角色授权，权限解析函数单测覆盖。")

    add_h3(doc, "Task 2.1：Dify 客户端 + httpx 单例")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/dify/client.py（DifyClient）", "BE-2", "list_apps / upload_file / chat_messages_stream"],
            ["app/core/lifespan.py（httpx + DifyClient 单例）", "BE-2", "worker 启动时初始化"],
            ["app/main.py 接入 lifespan", "BE-2", "lifespan 注入到 app.state.dify"],
            ["test_dify_client.py", "BE-2", "2 个测试通过"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 2.2：Agent 同步")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/models/{app,dify_api_key}.py", "BE-2", "apps + dify_api_keys 表"],
            ["app/apps/sync.py（upsert + API Key 加密入库）", "BE-2", "sync_apps_from_dify 可用"],
            ["app/schemas/app.py", "BE-2", "AppOut DTO"],
            ["app/apps/router.py（GET /api/apps/me MVP 占位）", "BE-2", "返回全量启用 App"],
            ["test_app_sync.py", "BE-2", "Mock Dify 同步测试通过"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 2.3：三态主体授权（user / dept / role）")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/models/app_authorization.py", "BE-1", "复合 PK + CheckConstraint"],
            ["app/apps/authorization.py", "BE-1", "user_can_access_app 解析函数"],
            ["test_authorization.py（5 个核心场景）", "BE-1", "user / dept / role / 复合 / 无授权"],
            ["/api/apps/me 升级为按用户解析可见列表", "BE-1", "权限过滤生效"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 2 准入标准")
    add_bullet(doc, "启动时自动从 Dify 拉取应用 + 加密存 API Key；")
    add_bullet(doc, "/api/apps/me 仅返回当前用户有权的 Agent；")
    add_bullet(doc, "user_can_access_app 5 个场景单测全过；")
    add_bullet(doc, "FastAPI 侧二次权限校验生效。")

    add_h2(doc, "4.5 阶段 3：对话代理 + SSE + 对话镜像（2 周）")
    add_body(doc, "目标：API 调通可流式对话，本地落库，对话可查询；并完成加密导出后端。")

    add_h3(doc, "Task 3.1：对话与消息数据模型")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/models/{conversation,message,audit_log}.py", "BE-1", "conversations(UUID) + messages + audit_logs"],
            ["app/schemas/conversation.py + message.py", "BE-1", "DTO"],
            ["alembic 迁移（dedupe_key UNIQUE）", "BE-1", "数据库就绪"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 3.2：SSE 流式对话代理")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/chat/service.py（user_msg 落库 + dify_payload 构造）", "BE-2", "权限校验 + 落库"],
            ["app/chat/router.py（POST /api/chat/send）", "BE-2", "StreamingResponse 透传"],
            ["_dify_event_generator 实现", "BE-2", "aiter_lines 逐 chunk 透传 + finally 写库"],
            ["超时 + 错误兜底（TimeoutException / Exception）", "BE-2", "event:error + 审计"],
            ["test_chat_sse.py（fake generator）", "BE-2", "透传 + finally 落库测试通过"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 3.3：对话历史查询 + 限流最小集")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/conversations/router.py（list / get / 分页）", "BE-1", "分页 + 软删过滤"],
            ["app/conversations/queries.py（iter_messages）", "BE-1", "流式分批 100 条"],
            ["app/core/rate_limit.py（内存令牌桶 MVP 版）", "BE-1", "每用户 60 req/min"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 3.4：加密导出后端端点")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/chat/exports.py（GET /api/conversations/{id}/export）", "BE-2", "流式 json/md + 审计三态"],
            ["app/conversations/batch_export.py（POST /api/conversations/export-batch）", "BE-2", "zipstream + 批量审计"],
            ["app/admin/audit_export.py（POST /api/admin/audit-logs/export）", "BE-2", "CSV 注入防护 + user_agent 字段"],
            ["app/files/downloads.py（GET /api/files/{file_id}/download）", "BE-2", "sanitize_filename"],
            ["app/utils/filename.py（sanitize_filename）", "BE-2", "路径穿越 / CRLF / bidi"],
            ["app/audit/service.py（audit_log 统一签名）", "BE-1", "所有导出端点共用"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 3.5：Dify Webhook 兜底")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/internal/webhook/dify.py（app-events + workflow-completed）", "BE-2", "Webhook 端点"],
            ["messages dedupe_key 唯一性保证", "BE-2", "INSERT ON CONFLICT DO NOTHING"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 3 准入标准")
    add_bullet(doc, "curl 调 /api/chat/send 看到 SSE 流，首字节延迟 ≤ 300ms；")
    add_bullet(doc, "对话完成后 messages 表同时有 user + assistant 两条；")
    add_bullet(doc, "客户端断流测试：assistant 消息仍落库 + 审计 completed=false；")
    add_bullet(doc, "Dify 5xx 时前端看到 event:error + 审计落库；")
    add_bullet(doc, "4 个导出端点 E2E 通过（含客户端断开场景）；")
    add_bullet(doc, "CSV 注入测试通过（action=cmd\\|… → 前缀 '）。")

    add_h2(doc, "4.6 阶段 4：文件上传（0.5 周）")
    add_body(doc, "目标：文件上传全链路可用，大小 / MIME / 文件名三重校验到位。")

    add_h3(doc, "Task 4.1：文件上传端点")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["app/files/upload.py（POST /api/chat/files）", "BE-2", "multipart 接收"],
            ["Content-Length + 流式双层校验", "BE-2", "20MB 上限"],
            ["MIME 白名单校验（pdf/docx/txt/md/png/jpeg）", "BE-2", "白名单命中"],
            ["secure_filename 清洗", "BE-2", "路径穿越 / CRLF 防御"],
            ["存本地 + 转 Dify /files/upload", "BE-2", "拿 dify_file_id"],
            ["上传审计（user_id / file_name / size / mime / sha256）", "BE-2", "audit_logs 落库"],
            ["test_file_upload.py（4 个异常场景）", "BE-2", "超限 / 非法 MIME / 路径穿越 / 正常"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 4 准入标准")
    add_bullet(doc, "20MB PDF 上传成功；")
    add_bullet(doc, "20MB+1B 文件返回 413；")
    add_bullet(doc, "非白名单 MIME（如 application/x-msdownload）返回 415；")
    add_bullet(doc, "恶意文件名（../../etc/passwd）被清洗；")
    add_bullet(doc, "上传记录出现在 audit_logs。")

    add_h2(doc, "4.7 阶段 5：前端骨架 + 登录 + 对话 UI（2 周）")
    add_body(doc, "目标：员工端 MVP 可用，可登录、可选 Agent、可流式对话、可看历史、可加密导出。")

    add_h3(doc, "Task 5.1：前端骨架")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["vite + ts + antd + zustand + axios 初始化", "FE-1", "脚手架就绪"],
            ["路由 + 布局（Login / Chat / History / Admin）", "FE-1", "可访问 /login /chat /history /admin"],
            ["axios 封装（credentials: include）", "FE-1", "cookie 自动带"],
            ["全局 store（auth / apps / 当前对话）", "FE-1", "Zustand store 3 个"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 5.2：登录页")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["Login.tsx（邮箱 + 密码 + 错误提示）", "FE-1", "登录成功后跳 /chat"],
            ["401 自动跳 /login", "FE-1", "access 过期处理"],
            ["refresh 自动续期（拦截器）", "FE-1", "用户无感"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 5.3：对话页（核心 UI）")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["Agent 切换 + 列表", "FE-1", "从 /api/apps/me 拉"],
            ["fetch + ReadableStream 消费 SSE", "FE-1", "流式渲染（不用 EventSource）"],
            ["消息列表 + 输入框 + 发送按钮", "FE-1", "基本对话 UI"],
            ["agent_done 事件后停止 loading", "FE-1", "UX 闭环"],
            ["文件上传按钮 + 进度", "FE-1", "可发文件给 Dify"],
            ["单会话导出按钮（带加密 checkbox）", "FE-1", "调导出端点"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 5.4：历史页 + 加密导出 UI")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["History.tsx（分页 + 列表）", "FE-1", "可看历史对话"],
            ["批量多选 + 导出按钮", "FE-1", "调 /api/conversations/export-batch"],
            ["<DownloadButton> + <PasswordDialog>", "FE-1", "统一下载入口"],
            ["useEncryptedDownload hook（PBKDF2 + envelope）", "FE-1", "加密下载流式"],
            ["/decrypt 公共页面", "FE-1", "envelope allow-list + 离线解密"],
            ["zxcvbn 强度评估", "FE-1", "弱密码红字警告"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 5 准入标准")
    add_bullet(doc, "登录 → /chat 看到 Agent 列表（仅可见）；")
    add_bullet(doc, "选 Agent → 输消息 → 流式看到回复；")
    add_bullet(doc, "历史页能看到刚才的对话；")
    add_bullet(doc, "选对话 → 加密下载 → 输密码 → 拿到 .enc；")
    add_bullet(doc, "/decrypt 上传 .enc + 同密码 → 拿到原文；")
    add_bullet(doc, "弱密码（zxcvbn=1）弹红字警告；")
    add_bullet(doc, "Chrome / Edge / Firefox 三浏览器均验证通过。")

    add_h2(doc, "4.8 阶段 6：管理员后台（1 周）")
    add_body(doc, "目标：PLATFORM_ADMIN 可管用户 / Agent 授权；APP_ADMIN 可看自己有权 Agent 的统计。")

    add_h3(doc, "Task 6.1：管理员路由聚合")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["/api/admin/users（CRUD）", "BE-1", "用户的 CRUD"],
            ["/api/admin/departments 补全 CRUD", "BE-1", "增删改查"],
            ["/api/admin/roles 补全 CRUD", "BE-1", "增删改查"],
            ["/api/admin/apps/{id}/authorizations", "BE-1", "三态授权 CRUD"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 6.2：前端管理后台")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["Admin/Users.tsx（列表 + 新建 / 编辑 / 禁用）", "FE-1", "可管用户"],
            ["Admin/Depts.tsx（树形 + CRUD）", "FE-1", "可管部门"],
            ["Admin/Roles.tsx（列表 + CRUD）", "FE-1", "可管角色"],
            ["Admin/AgentAuthorizations.tsx", "FE-1", "选 Agent → 选主体类型 → 授权 / 撤销"],
            ["Admin/AuditLogs.tsx（按时间 + 动作查询 + 导出）", "FE-1", "审计可查可导出"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 6 准入标准")
    add_bullet(doc, "PLATFORM_ADMIN 可创建用户、分配角色、加入部门；")
    add_bullet(doc, "PLATFORM_ADMIN 可给 Agent 授权给指定用户 / 部门 / 角色；")
    add_bullet(doc, "被授权用户立即在 /api/apps/me 看到该 Agent；")
    add_bullet(doc, "审计页面可按时间窗口 + 动作类型查询；")
    add_bullet(doc, "审计导出 CSV 含 user_agent 列 + 公式注入防护。")

    add_h2(doc, "4.9 阶段 7：生产部署 + 可观测性（1 周）")
    add_body(doc, "目标：生产环境可用，监控告警最小集就位。")

    add_h3(doc, "Task 7.1：生产部署编排")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["生产 docker-compose.yml（含 2 副本 FastAPI + Nginx）", "SRE", "docker stack deploy"],
            ["Nginx 配置（proxy_buffering off + /decrypt CSP）", "SRE", "SSE + CSP 正确"],
            ["HTTPS 证书（Nginx + Let's Encrypt / 公司 CA）", "SRE", "TLS 1.2+"],
            ["健康检查 + 启动顺序脚本", "SRE", "起停顺序文档化"],
            ["Postgres 备份脚本（每日 cron）", "SRE", "保留 7 天"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 7.2：可观测性最小集")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["loguru JSON formatter + 关键字段", "BE-1", "结构化日志"],
            ["关键节点日志：JWT 鉴权失败 / 权限拒绝 / SSE 流开始结束 / Dify 调用耗时", "BE-1", "可 grep"],
            ["worker_id 注入（lifespan）", "BE-1", "排障可定位 worker"],
            ["SSE 首字节延迟打点（中间件）", "BE-1", "P50 / P95 可统计"],
            ["告警脚本（脚本级，MVP 阶段不接 Prometheus）", "SRE", "Postgres 表过大 / 错误率超阈邮件告警"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "Task 7.3：上线演练 + 验收")
    add_data_table(doc,
        headers=["工作项", "负责人", "产出"],
        rows=[
            ["用户验收（业务代表走 5 个核心场景）", "PM + 业务", "UAT 报告"],
            ["E2E 全量回归（Playwright）", "测试", "全过"],
            ["性能压测（对话并发 200）", "SRE", "P95 延迟报告"],
            ["安全 review 收尾报告", "安全", "签发上线许可"],
            ["上线 + 7×24 值守表（首周）", "SRE", "故障响应"],
        ],
        col_widths=[6.5, 3.0, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "阶段 7 准入标准")
    add_bullet(doc, "生产环境 docker stack deploy 一次成功；")
    add_bullet(doc, "HTTPS 可访问，SSE 首字节延迟 P95 ≤ 500ms；")
    add_bullet(doc, "200 路对话并发稳定；")
    add_bullet(doc, "关键告警脚本验证通过；")
    add_bullet(doc, "UAT 通过 + 安全 review 通过。")

    add_page_break(doc)


def write_chapter_5(doc):
    add_h1(doc, "第五章  任务依赖与关键路径")

    add_h2(doc, "5.1 任务依赖关系")
    add_body(doc, "本节梳理关键任务之间的依赖关系，避免出现\"下游任务比上游更早开始\"的情况。")
    add_data_table(doc,
        headers=["上游任务", "下游任务", "依赖内容"],
        rows=[
            ["Task 0.1 仓库骨架", "Task 0.2 Dify 部署", "目录结构 + docker-compose.yml"],
            ["Task 0.2 Dify 部署", "Task 2.1 Dify 客户端", "Dify 端点可达 + API Key"],
            ["Task 1.1 DB 配置 + 迁移", "Task 1.2 JWT", "Base 模型 + get_db"],
            ["Task 1.2 JWT", "Task 1.3 认证 API", "签发 / 校验函数"],
            ["Task 1.3 认证 API", "Task 1.4 组织架构", "current_user 依赖"],
            ["Task 1.4 组织架构", "Task 2.3 授权", "User / Department / Role 表"],
            ["Task 2.1 Dify 客户端", "Task 2.2 Agent 同步", "DifyClient.list_apps"],
            ["Task 2.2 Agent 同步 + Task 2.3 授权", "Task 3.2 对话代理", "App 模型 + 权限解析"],
            ["Task 3.1 对话数据模型", "Task 3.2 SSE + Task 3.3 历史查询", "conversations / messages 表"],
            ["Task 3.2 SSE 代理", "Task 3.4 加密导出后端", "对话镜像落库"],
            ["Task 3.4 加密导出后端", "Task 5.4 加密导出前端", "后端端点 + envelope 格式"],
            ["Task 4.1 文件上传", "Task 5.3 对话页文件按钮", "/api/chat/files 端点"],
            ["Task 1.x 后端", "Task 5.x 前端", "后端 API 稳定后可联调"],
            ["Task 6.x 管理后台", "Task 7.x 上线", "管理功能就绪才能上线"],
        ],
        col_widths=[4.0, 4.0, 7.3],
        first_col_bold=True,
    )

    add_h2(doc, "5.2 关键路径分析")
    add_body(doc, "项目的关键路径（Critical Path）如下：")
    add_quote(doc,
        "0.1 → 0.2 → 1.1 → 1.2 → 1.3 → 1.4 → 2.1 → 2.2 → 2.3 → 3.1 → 3.2 → 3.3 → 3.4 → 4.1 → 5.x → 6.x → 7.x\n\n"
        "关键路径总工期：1 + 1 + 1.5 + 2 + 0.5 + 2 + 1 + 1 = 10 周\n\n"
        "关键路径上的任何延期都会直接影响项目上线时间，必须重点保障资源投入。")

    add_h3(doc, "5.2.1 关键路径风险点")
    add_data_table(doc,
        headers=["节点", "风险", "影响"],
        rows=[
            ["Task 0.2 Dify 部署", "镜像版本、Postgres 外部化、网络打通", "拖 1 周 = 整体 +1 周"],
            ["Task 3.2 SSE 三层联调", "nginx 缓冲 / httpx 流 / 前端 fetch 协调", "拖 1.5 周 = 整体 +1.5 周"],
            ["Task 3.4 加密导出后端", "审计三态 + 流式 zip + CSV 注入", "拖 1 周 = 整体 +1 周"],
            ["Task 5.x 前端联调", "后端 API 稳定 + 前端 SSE 消费", "拖 1 周 = 整体 +1 周"],
        ],
        col_widths=[4.0, 7.0, 4.3],
        first_col_bold=True,
    )

    add_h2(doc, "5.3 并行执行机会")
    add_body(doc, "以下任务可并行执行，最大化资源利用率：")
    add_data_table(doc,
        headers=["并行组 A", "并行组 B", "并行依据"],
        rows=[
            ["Task 0.2 Dify 部署", "Task 1.1 后端骨架",
             "SRE vs BE-1，无共同依赖"],
            ["Task 1.2 JWT", "Task 1.4 组织架构",
             "BE-1 同一开发可串行；BE-2 可并行做 Dify 适配"],
            ["Task 3.2 SSE", "Task 3.3 历史查询",
             "BE-1 与 BE-2 并行"],
            ["Task 3.4 加密导出后端", "Task 4.1 文件上传",
             "BE-2 内部串行；与 Task 5.x 前端并行"],
            ["Task 5.x 前端", "Task 6.x 管理后台",
             "FE-1 在阶段 5 完成后切到阶段 6；可与阶段 5 收尾并行"],
            ["Task 7.1 部署编排", "Task 7.2 可观测性",
             "SRE 与 BE-1 并行"],
        ],
        col_widths=[5.0, 5.0, 5.3],
        first_col_bold=True,
    )

    add_quote(doc,
        "并行机会通常可压缩 1–2 周，但会增加集成风险。"
        "建议在阶段 3 后半段（前端联调）开始前预留 2–3 天的串行集成窗口，"
        "避免并行产物接口不一致导致返工。")

    add_page_break(doc)


def write_chapter_6(doc):
    add_h1(doc, "第六章  里程碑与交付物")

    add_h2(doc, "6.1 里程碑表")
    add_data_table(doc,
        headers=["里程碑", "时间节点", "标志事件", "准入检查"],
        rows=[
            ["M1 - 环境就绪", "第 1 周末（W1）",
             "Postgres + Dify 起 + alembic 迁移成功 + 健康检查通过",
             "curl 调通 Dify API + 后端 /api/health 200"],
            ["M2 - 鉴权可用", "第 2.5 周末（W2.5）",
             "登录 + JWT + CSRF 单测通过 + E2E 跑通",
             "pytest 100% + 跨域 POST 403"],
            ["M3 - 授权生效", "第 3.5 周末（W3.5）",
             "三态授权解析 + 权限校验单测覆盖",
             "5 个授权场景单测全过"],
            ["M4 - 对话可用", "第 5.5 周末（W5.5）",
             "SSE 流式对话 + 落库 + finally 兜底",
             "首字节 ≤ 300ms + 客户端断流兜底"],
            ["M5 - 文件上传", "第 6 周末（W6）",
             "20MB + MIME 白名单 + 文件名清洗全链路",
             "4 个异常场景单测全过"],
            ["M6 - 前端 MVP", "第 8 周末（W8）",
             "用户可登录 + 对话 + 历史 + 加密导出",
             "3 浏览器 E2E 全过"],
            ["M7 - 管理后台", "第 9 周末（W9）",
             "PLATFORM_ADMIN 可管用户 / Agent 授权",
             "UAT 通过"],
            ["M8 - 上线", "第 10 周末（W10）",
             "生产部署 + 监控告警就位 + 验收通过",
             "压测 + 安全 review + UAT 全过"],
        ],
        col_widths=[3.0, 2.5, 5.5, 4.3],
        first_col_bold=True,
    )

    add_h2(doc, "6.2 交付物清单")

    add_h3(doc, "6.2.1 阶段 0 交付物")
    add_bullet(doc, "git 仓库（含 backend/ frontend/ deploy/ docs/ 目录结构）")
    add_bullet(doc, "docker-compose.yml + nginx/conf.d/portal.conf")
    add_bullet(doc, "Dify 1.1.0 部署完成（PG / Redis 外部化）")
    add_bullet(doc, "DIFY_DEPLOY.md（部署验证记录）")
    add_bullet(doc, "README.md（团队上手文档）")

    add_h3(doc, "6.2.2 阶段 1 交付物")
    add_bullet(doc, "config.py + db/{base,session,deps}.py + models 5 张表")
    add_bullet(doc, "security.py（bcrypt + JWT + refresh hash）")
    add_bullet(doc, "auth/{router,service,deps}.py + middleware.py（CSRF）")
    add_bullet(doc, "admin/{users,depts,roles} 基础 CRUD")
    add_bullet(doc, "alembic 迁移 + pytest 套件通过")

    add_h3(doc, "6.2.3 阶段 2 交付物")
    add_bullet(doc, "dify/client.py + core/lifespan.py（httpx + DifyClient 单例）")
    add_bullet(doc, "apps/sync.py（启动时同步 + 加密 API Key 入库）")
    add_bullet(doc, "apps/authorization.py（user_can_access_app 解析函数）")
    add_bullet(doc, "单测覆盖：5 个授权场景")

    add_h3(doc, "6.2.4 阶段 3 交付物")
    add_bullet(doc, "models/{conversation,message,audit_log}.py")
    add_bullet(doc, "chat/router.py + service.py（SSE 透传 + 错误兜底 + finally 写库）")
    add_bullet(doc, "conversations/router.py + queries.py（分页 + 流式查询）")
    add_bullet(doc, "4 个导出端点（单/批量/审计/附件）+ sanitize_filename + 审计三态")
    add_bullet(doc, "internal/webhook/dify.py + messages dedupe_key UNIQUE")

    add_h3(doc, "6.2.5 阶段 4 交付物")
    add_bullet(doc, "files/upload.py（Content-Length + 流式 + MIME + 文件名 + 审计）")
    add_bullet(doc, "test_file_upload.py（4 个异常场景）")

    add_h3(doc, "6.2.6 阶段 5 交付物")
    add_bullet(doc, "前端骨架（路由 + 布局 + store + axios）")
    add_bullet(doc, "Login / Chat / History / Decrypt 4 个页面")
    add_bullet(doc, "lib/crypto/ 6 个模块（envelope / deriveKey / encryptStream / decryptFile / sanitizeFilename）")
    add_bullet(doc, "hooks/useEncryptedDownload.ts + components/{DownloadButton,PasswordDialog}.tsx")
    add_bullet(doc, "Vitest 单测套件（envelope / deriveKey / sanitize 等）")

    add_h3(doc, "6.2.7 阶段 6 交付物")
    add_bullet(doc, "后端 /api/admin/* 完整 CRUD")
    add_bullet(doc, "前端 Admin/{Users,Depts,Roles,AgentAuthorizations,AuditLogs}.tsx")

    add_h3(doc, "6.2.8 阶段 7 交付物")
    add_bullet(doc, "生产 docker-compose.yml + Nginx（含 CSP）+ HTTPS 证书")
    add_bullet(doc, "loguru JSON 日志 + 关键节点打点 + worker_id")
    add_bullet(doc, "Postgres 备份脚本 + 告警脚本")
    add_bullet(doc, "UAT 报告 + 压测报告 + 安全 review 报告 + 上线签发单")

    add_page_break(doc)


def write_chapter_7(doc):
    add_h1(doc, "第七章  资源与预算")

    add_h2(doc, "7.1 人力资源")
    add_data_table(doc,
        headers=["角色", "人数", "人周", "主要工作时段"],
        rows=[
            ["项目经理", "1", "10", "全程"],
            ["后端工程师 BE-1（主程）", "1", "10", "全程"],
            ["后端工程师 BE-2（Dify / SSE / 文件）", "1", "10", "全程"],
            ["前端工程师 FE-1", "1", "10", "全程（前期兼职，后期全职）"],
            ["SRE / DevOps", "0.3", "3", "阶段 0（1 周）+ 阶段 7（1 周）+ 持续 0.1 兼职"],
            ["安全 review", "0.2", "2", "阶段 1（0.5 周）+ 阶段 3（0.5 周）+ 阶段 4（0.3 周）+ 阶段 7（0.5 周）"],
            ["测试工程师（兼）", "0.5", "5", "阶段 3 末（1 周）+ 阶段 5（1 周）+ 阶段 6（1 周）+ 阶段 7（1 周）"],
            ["业务代表（产品）", "0.3", "3", "全程"],
            ["合计", "—", "约 53 人周", "—"],
        ],
        col_widths=[4.5, 1.5, 1.5, 7.8],
        first_col_bold=True,
    )

    add_h2(doc, "7.2 硬件资源")
    add_data_table(doc,
        headers=["资源", "规格", "数量", "用途"],
        rows=[
            ["开发机", "16C32G / 1TB SSD", "4", "后端 / 前端 / SRE 开发"],
            ["测试服务器", "8C16G / 500G SSD", "1", "集成测试 + 性能压测"],
            ["生产服务器 - FastAPI 节点", "2C4G / 50G SSD", "2", "FastAPI 容器（2 实例）"],
            ["生产服务器 - Dify 节点", "8C16G / 200G SSD", "1", "Dify 社区版（5+ 微服务）"],
            ["生产服务器 - Postgres 节点", "4C8G / 500G SSD", "1", "主数据库"],
            ["生产服务器 - Nginx 节点", "2C4G / 50G SSD", "1", "反代 + 静态服务"],
            ["备份存储", "1TB NAS", "1", "Postgres 每日备份"],
        ],
        col_widths=[5.0, 3.5, 1.5, 5.3],
        first_col_bold=True,
    )

    add_h2(doc, "7.3 第三方服务与软件")
    add_data_table(doc,
        headers=["项目", "规格 / 数量", "费用"],
        rows=[
            ["Dify 社区版（自托管）", "1.1.0", "免费（Apache 2.0）"],
            ["PostgreSQL", "16-alpine", "免费（PostgreSQL License）"],
            ["Redis（仅 Dify 用）", "7-alpine", "免费"],
            ["Nginx", "1.25-alpine", "免费"],
            ["React / FastAPI / SQLAlchemy 等", "开源", "免费"],
            ["Let's Encrypt 证书", "DV 证书", "免费（如用公司 CA 则 0）"],
            ["zxcvbn（密码强度）", "DropBox 开源", "免费"],
            ["Penumbra（浏览器加密）", "@transcend-io/penumbra 8.1.4", "免费（Apache 2.0）"],
            ["1Password / 类似密码管理", "团队版", "已有预算"],
            ["合计", "—", "除 1Password 外 0 软件成本"],
        ],
        col_widths=[5.0, 5.0, 5.3],
        first_col_bold=True,
    )

    add_quote(doc,
        "MVP 阶段软件成本接近 0，主要成本是人力 + 硬件 + 时间。\n"
        "硬件：开发机 + 测试机 + 4 台生产服务器，可借用公司现有资源。")

    add_page_break(doc)


def write_chapter_8(doc):
    add_h1(doc, "第八章  风险与应对")

    add_h2(doc, "8.1 风险登记册")

    add_h3(doc, "8.1.1 技术风险")
    add_data_table(doc,
        headers=["#", "风险", "严重度", "概率", "影响", "缓解措施", "责任人"],
        rows=[
            ["T-01", "SSE 缓冲导致体感差", "高", "中", "首字节延迟 > 1s",
             "header + 代码双重保险（X-Accel-Buffering:no + proxy_buffering off + aiter_lines）；上线前压测",
             "BE-2"],
            ["T-02", "Dify API 变更破坏适配层", "中", "低", "核心功能失效",
             "锁版本（langgenius/dify-api:1.1.0 digest），升级走专项 review",
             "BE-2"],
            ["T-03", "对话数据双写不一致", "中", "中", "消息丢失 / 重复",
             "Webhook 兜底 + 定期对账脚本（仅告警不删）+ messages.dedupe_key UNIQUE",
             "BE-1"],
            ["T-04", "LLM 成本失控", "中", "高", "预算超支",
             "MVP 阶段不加配额；先观察一周；二期按部门配额",
             "PM + 业务"],
            ["T-05", "前端 SSE 断线导致内容丢失", "中", "中", "用户体验差",
             "MVP 不做断点续传；二期实现 fetch 手动中断 + 重发；当前 finally 块保证 assistant 消息必落库",
             "FE-1"],
            ["T-06", "Dify 部署 / 升级影响生产", "高", "低", "Dify 不可用",
             "Dify 独立部署（与本项目解耦）；升级前备份 Dify DB；滚动升级",
             "SRE"],
        ],
        col_widths=[0.8, 3.5, 1.2, 1.0, 2.5, 4.5, 1.8],
        first_col_bold=True,
    )

    add_h3(doc, "8.1.2 进度风险")
    add_data_table(doc,
        headers=["#", "风险", "严重度", "概率", "影响", "缓解措施", "责任人"],
        rows=[
            ["S-01", "Dify 部署调通延期", "高", "中", "+1 周",
             "阶段 0 提前做部署验证；外部化 PG / Redis 优先；镜像版本锁定",
             "SRE"],
            ["S-02", "SSE 三层联调调试周期长", "中", "高", "+1.5 周",
             "nginx + uvicorn + httpx 三层同时调；预置调试工具（worker_id 日志）；压测脚本提前准备",
             "BE-2"],
            ["S-03", "多态授权一致性单测覆盖不全", "中", "中", "Bug 漏到生产",
             "强制 user_can_access_app 100% 覆盖率；复合授权单测",
             "BE-1"],
            ["S-04", "加密导出浏览器兼容性问题", "中", "中", "用户报障",
             "目标 Chrome / Edge / Firefox ≥ 102；明确不支持 Safari；上线前跨浏览器验证",
             "FE-1"],
            ["S-05", "关键人员离职 / 缺席", "高", "低", "+2 周以上",
             "代码 review 机制；双人关键模块；文档同步更新",
             "PM"],
        ],
        col_widths=[0.8, 3.5, 1.2, 1.0, 2.5, 4.5, 1.8],
        first_col_bold=True,
    )

    add_h3(doc, "8.1.3 安全风险")
    add_data_table(doc,
        headers=["#", "风险", "严重度", "概率", "影响", "缓解措施", "责任人"],
        rows=[
            ["X-01", "离职员工用旧 cookie 调通", "中", "中", "越权访问",
             "JWT 短过期（15min）+ jti 黑名单（二期）；refresh 主动撤销；离职触发软删对话",
             "安全"],
            ["X-02", "CSRF / 文件上传安全漏洞", "高", "低", "安全事件",
             "SameSite=Strict + Origin 校验 + 文件白名单 + 大小限制；E2E 覆盖异常路径",
             "BE-1"],
            ["X-03", "SECRET_KEY 单点泄漏", "中", "低", "全平台密钥风险",
             "已拆分 JWT_SECRET + ENCRYPTION_KEY；独立轮转；KMS 管理",
             "SRE"],
            ["X-04", "弱密码导致 .enc 可破解", "中", "高", "用户数据泄露",
             "zxcvbn 警告 + 12 字符最小 + 2 类字符；员工遗忘密码风险已在 /decrypt 提示",
             "FE-1"],
            ["X-05", "CSV 公式注入", "中", "中", "管理员被攻击",
             "Python csv 模块 + sanitize_csv_cell 前缀",
             "BE-2"],
            ["X-06", "origName 文件名注入", "低", "中", "下载文件被劫持",
             "sanitize_filename 服务端 + 前端双重清洗；防路径穿越 / CRLF / bidi",
             "BE-2"],
            ["X-07", "envelope 字段值未校验", "中", "低", "解密页被攻击",
             "/decrypt 强制 allow-list（alg / kdf / iter 范围 / ivLen / version）",
             "FE-1"],
            ["X-08", "审计日志泄露用户下载内容", "低", "低", "敏感信息扩散",
             "audit_logs 只存元数据，不存原文",
             "BE-1"],
        ],
        col_widths=[0.8, 3.5, 1.2, 1.0, 2.5, 4.5, 1.8],
        first_col_bold=True,
    )

    add_h2(doc, "8.2 风险监控与升级")
    add_data_table(doc,
        headers=["风险等级", "触发条件", "响应时间", "响应人"],
        rows=[
            ["高", "达到发生概率阈值 + 不可接受影响", "4 小时内", "PM 立即升级到 CCB"],
            ["中", "已发生或有明确发生信号", "24 小时内", "PM + 责任人 + 业务代表"],
            ["低", "有发生可能性但暂未发生", "周评审会同步", "责任人按计划监控"],
    ],
        col_widths=[2.0, 6.0, 3.0, 4.3],
        first_col_bold=True,
    )

    add_quote(doc,
        "风险升级机制：\n"
        "- 风险等级\"中\"及以上，每日 stand-up 必同步；\n"
        "- 风险等级\"高\"且即将发生，立即召集相关方专题会；\n"
        "- 风险等级\"高\"已发生，PM 在 4 小时内升级到 CCB + 业务代表。")

    add_page_break(doc)


def write_chapter_9(doc):
    add_h1(doc, "第九章  质量保证")

    add_h2(doc, "9.1 质量目标")
    add_data_table(doc,
        headers=["质量维度", "目标值", "度量方法"],
        rows=[
            ["功能正确性", "P0 / P1 缺陷 0", "E2E 全过 + UAT 通过"],
            ["后端单测覆盖率", "关键模块 ≥ 70%，鉴权 / 授权 / SSE 100%",
             "pytest --cov=app --cov-report=term-missing"],
            ["前端单测覆盖率", "lib/crypto ≥ 90%",
             "vitest --coverage"],
            ["E2E 关键路径", "8 个核心场景全过",
             "Playwright"],
            ["首字节延迟", "P95 ≤ 500ms", "压测 + 生产监控"],
            ["可用性", "≥ 99.5%", "上线后统计"],
            ["安全漏洞", "上线前安全 review 0 高危", "安全 review 报告"],
        ],
        col_widths=[4.0, 5.5, 5.8],
        first_col_bold=True,
    )

    add_h2(doc, "9.2 评审机制")
    add_data_table(doc,
        headers=["评审类型", "时机", "参与者", "产出"],
        rows=[
            ["代码 review（PR）", "每个 PR 合并前", "至少 1 名主程 + 1 名其他",
             "PR 评审通过 + CI 通过"],
            ["阶段准入评审", "每个阶段末", "PM + 主程 + 业务代表",
             "阶段交付物验收报告"],
            ["架构评审", "阶段 1 / 3 / 5 末", "主程 + 外部架构师（如有）",
             "架构决策记录（ADR）"],
            ["安全 review", "阶段 1 / 3 / 4 / 7 关键节点", "安全团队",
             "安全评估报告 + 风险清单"],
            ["UAT 验收", "阶段 7 末", "业务代表 + 真实用户抽样",
             "UAT 报告 + 签发上线许可"],
        ],
        col_widths=[3.0, 3.0, 4.0, 5.3],
        first_col_bold=True,
    )

    add_h2(doc, "9.3 测试策略")
    add_data_table(doc,
        headers=["层级", "范围", "工具", "准入门槛"],
        rows=[
            ["单元测试", "app/* 业务模块、lib/crypto/*",
             "pytest + Vitest", "关键模块覆盖率 ≥ 70%"],
            ["集成测试", "所有 API 端点（流式 + 异常路径）",
             "pytest + httpx", "E2E 跑通前必须全过"],
            ["E2E", "8 个核心业务路径",
             "Playwright", "上线前 100% 通过"],
            ["手动验证", "首字节延迟、断流兜底、3 浏览器兼容",
             "人工", "SRE + FE-1 联合作业"],
            ["性能压测", "200 路对话并发",
             "locust / wrk", "P95 ≤ 500ms"],
            ["安全扫描", "依赖漏洞 + 密钥泄漏",
             "pip-audit / npm audit / git-secrets",
             "0 高危"],
        ],
        col_widths=[2.5, 5.0, 3.0, 4.8],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_chapter_10(doc):
    add_h1(doc, "第十章  变更管理")

    add_h2(doc, "10.1 变更流程")
    add_code_block(doc, """
需求变更提出
   │
   ▼
PM 初步评估（影响：进度 / 范围 / 成本 / 质量）
   │
   ├─── 小变更（不影响 MVP 准入） ──→ PM 决策 → 同步到周评审
   │
   ├─── 中变更（影响 1-2 周进度） ──→ CCB 评审 → 决策
   │
   └─── 大变更（影响 1 周以上 / 范围 / 成本） ──→ CCB + 业务方 → 决策

所有变更 → 写入变更日志 → 更新本执行计划 + 技术方案框架
    """, language="变更流程")

    add_h3(doc, "10.1.1 变更日志模板")
    add_data_table(doc,
        headers=["字段", "示例"],
        rows=[
            ["变更编号", "CR-2026-009"],
            ["提出日期", "2026-09-15"],
            ["提出人", "业务代表 王五"],
            ["变更描述", "在 MVP 中增加 SSO 接口预留"],
            ["影响分析", "进度 +0.5 天 / 范围 +10 行 / 成本 0 / 质量 +"],
            ["决策", "批准（CCB 2026-09-16）"],
            ["实施人", "BE-1"],
            ["完成日期", "2026-09-17"],
            ["状态", "已实施"],
        ],
        col_widths=[3.5, 11.8],
        first_col_bold=True,
    )

    add_h2(doc, "10.2 变更控制委员会（CCB）")
    add_data_table(doc,
        headers=["角色", "人员", "职责"],
        rows=[
            ["CCB 主席", "数字化办公室主任", "变更最终决策 + 升级处理"],
            ["常任委员", "应用开发组组长", "技术可行性 + 资源协调"],
            ["常任委员", "基础架构组组长", "基础设施影响评估"],
            ["常任委员", "安全团队代表", "安全影响评估"],
            ["常任委员", "PM", "变更材料准备 + 进度影响评估"],
            ["列席", "业务代表", "需求澄清（非决策权）"],
        ],
        col_widths=[3.0, 4.5, 7.8],
        first_col_bold=True,
    )

    add_quote(doc,
        "CCB 会议：按需召集（一般每周 1 次定时 + 紧急随时）。\n"
        "决策需 CCB 主席 + 半数常任委员在场；表决多数通过。\n"
        "紧急变更可走\"快速通道\"：PM + 主席 + 1 名相关常任委员口头批准，"
        "24 小时内补走正式流程。")

    add_page_break(doc)


def write_chapter_11(doc):
    add_h1(doc, "第十一章  沟通与汇报")

    add_h2(doc, "11.1 会议节奏")
    add_data_table(doc,
        headers=["会议", "频率", "时长", "参与", "输入 / 输出"],
        rows=[
            ["每日 stand-up", "工作日 09:30", "15 分钟", "全体开发",
             "昨日 / 今日 / 阻塞（口头）"],
            ["周计划会", "周一 10:00", "1 小时", "PM + 全体开发",
             "周计划文档 + 风险登记册更新"],
            ["周评审会", "周五 16:00", "1 小时", "全体 + 业务代表",
             "本周演示 + 验收清单"],
            ["阶段评审", "阶段末", "2 小时", "全体 + 干系人",
             "阶段交付物 + 阶段准入报告"],
            ["月度汇报", "每月最后一周五", "1 小时", "PM + 干系人",
             "月度状态报告（含趋势）"],
            ["临时专题会", "按需", "30-60 分钟", "相关人员",
             "专项问题（如安全评审、重大 Bug）"],
            ["CCB 会议", "按需", "1 小时", "CCB 委员",
             "变更决策 + 风险升级"],
        ],
        col_widths=[3.0, 2.5, 1.5, 3.5, 4.8],
        first_col_bold=True,
    )

    add_h2(doc, "11.2 报告机制")

    add_h3(doc, "11.2.1 日常同步")
    add_bullet(doc, "Slack / 钉钉项目群：阻塞 / 风险 / 决策 5 分钟内同步；")
    add_bullet(doc, "GitHub / GitLab Issues：任务状态实时更新；")
    add_bullet(doc, "Confluence / 内部 Wiki：本计划 + ADR + 变更日志统一管理。")

    add_h3(doc, "11.2.2 周报模板")
    add_data_table(doc,
        headers=["章节", "内容"],
        rows=[
            ["本周亮点", "完成的里程碑 / 重要交付"],
            ["下周计划", "下周要做的关键任务"],
            ["风险 / 阻塞", "已识别风险 + 升级项"],
            ["度量数据", "燃尽图 / 完成率 / 缺陷趋势"],
            ["变更", "本周通过的变更"],
            ["需要支持", "需要干系人支持的事项"],
        ],
        col_widths=[3.0, 12.3],
        first_col_bold=True,
    )

    add_h3(doc, "11.2.3 月度报告模板")
    add_data_table(doc,
        headers=["章节", "内容"],
        rows=[
            ["整体进度", "完成阶段 / 总进度百分比 / 与计划偏差"],
            ["里程碑达成", "已达成 / 即将达成 / 已延期的里程碑"],
            ["风险总览", "按等级分类 + 趋势"],
            ["变更总览", "本月变更数 + 影响"],
            ["资源使用", "人周 / 预算执行"],
            ["下月重点", "下月要解决的关键问题"],
        ],
        col_widths=[3.0, 12.3],
        first_col_bold=True,
    )

    add_h2(doc, "11.3 升级机制")
    add_data_table(doc,
        headers=["事项", "升级路径"],
        rows=[
            ["技术决策争议（BE-1 vs BE-2）", "主程仲裁 → 仍争议则 PM 决策"],
            ["范围 / 进度争议", "PM → CCB"],
            ["安全风险", "安全代表 → PM → CCB → 业务方"],
            ["关键人员缺席 / 离职", "PM 立即评估影响 → CCB"],
            ["业务方对交付物不满意", "PM 协调 → 仍不满意则业务方代表上报"],
            ["外部依赖延期", "PM 评估 → CCB（如影响 > 1 周）"],
        ],
        col_widths=[5.0, 10.3],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_appendix_a(doc):
    add_h1(doc, "附录 A  阶段 0–4 任务详细清单")

    add_body(doc, "本附录详列阶段 0–4 的所有任务，阶段 5–7 任务见正文章节 4.7–4.9。")

    add_h2(doc, "A.1 阶段 0 任务")
    add_data_table(doc,
        headers=["任务", "步骤", "产出"],
        rows=[
            ["Task 0.1.1", "git init + 配置 user", "git 仓库"],
            ["Task 0.1.2", "创建 .gitignore", ".gitignore"],
            ["Task 0.1.3", "backend/pyproject.toml", "可 pip install -e .[dev]"],
            ["Task 0.1.4", "backend/.env.example", "环境变量模板"],
            ["Task 0.1.5", "backend/app/main.py 最小", "/api/health 200"],
            ["Task 0.1.6", "frontend/package.json", "可 npm install"],
            ["Task 0.1.7", "frontend/vite.config.ts", "/api 代理配置"],
            ["Task 0.1.8", "frontend/tsconfig.json", "TS 编译通过"],
            ["Task 0.1.9", "frontend/index.html", "挂载 root"],
            ["Task 0.1.10", "deploy/docker-compose.yml", "Postgres 容器"],
            ["Task 0.1.11", "deploy/Makefile", "make up/down"],
            ["Task 0.1.12", "deploy/nginx/conf.d/portal.conf", "Nginx 配置"],
            ["Task 0.1.13", "README.md", "团队上手文档"],
            ["Task 0.2.1", "克隆 Dify 1.1.0", "Dify 仓库"],
            ["Task 0.2.2", "外部化 Postgres", "Dify 跑在外部 PG"],
            ["Task 0.2.3", "创建 dify 数据库", "DB 就绪"],
            ["Task 0.2.4", "启动 Dify compose", "5+ 容器 healthy"],
            ["Task 0.2.5", "初始化 Dify 管理员", "可登录 Dify Web"],
            ["Task 0.2.6", "验证 /v1/setup/initialize-status", "返回 finished:true"],
            ["Task 0.2.7", "创建测试 Agent", "test-agent + API Key"],
            ["Task 0.2.8", "验证应用 API 调通", "SSE 流可见"],
            ["Task 0.2.9", "DIFY_DEPLOY.md", "凭据入 1Password"],
        ],
        col_widths=[3.0, 7.0, 5.3],
        first_col_bold=True,
    )

    add_h2(doc, "A.2 阶段 1 任务")
    add_data_table(doc,
        headers=["任务", "步骤", "产出"],
        rows=[
            ["Task 1.1.1", "core/config.py（pydantic-settings）", "get_settings()"],
            ["Task 1.1.2", "db/base.py + session.py + deps.py", "async engine + get_db"],
            ["Task 1.1.3", "models/{user,dept,role,refresh_token}.py", "5 张表 ORM"],
            ["Task 1.1.4", "alembic init + env.py", "alembic 命令就绪"],
            ["Task 1.1.5", "alembic revision autogenerate", "首个迁移"],
            ["Task 1.1.6", "alembic upgrade head", "表创建成功"],
            ["Task 1.1.7", "tests/conftest.py", "session + engine fixture"],
            ["Task 1.2.1", "core/security.py", "hash + JWT + refresh hash"],
            ["Task 1.2.2", "test_security.py（3 用例）", "全过"],
            ["Task 1.3.1", "schemas/auth.py", "LoginRequest / MeResponse"],
            ["Task 1.3.2", "auth/service.py", "authenticate + issue + rotate"],
            ["Task 1.3.3", "auth/router.py", "login / logout / refresh / me"],
            ["Task 1.3.4", "auth/deps.py", "current_user 依赖"],
            ["Task 1.3.5", "core/middleware.py", "CSRFMiddleware"],
            ["Task 1.3.6", "main.py 接入", "CORS + CSRF + routers"],
            ["Task 1.3.7", "test_auth.py（6 用例）", "全过"],
            ["Task 1.4.1", "models/{dept,role}.py 完善", "物化路径 + CheckConstraint"],
            ["Task 1.4.2", "schemas/{dept,role}.py", "Create / Out DTO"],
            ["Task 1.4.3", "depts/router.py", "POST / GET /api/admin/departments"],
            ["Task 1.4.4", "roles/router.py", "POST / GET /api/admin/roles"],
            ["Task 1.4.5", "admin/router.py 聚合", "/api/admin/* 完整"],
            ["Task 1.4.6", "test_admin.py", "PLATFORM_ADMIN 验证"],
        ],
        col_widths=[3.0, 7.0, 5.3],
        first_col_bold=True,
    )

    add_h2(doc, "A.3 阶段 2 任务")
    add_data_table(doc,
        headers=["任务", "步骤", "产出"],
        rows=[
            ["Task 2.1.1", "dify/client.py", "DifyClient + 4 个方法"],
            ["Task 2.1.2", "core/lifespan.py", "httpx + DifyClient 单例"],
            ["Task 2.1.3", "main.py 接入 lifespan", "app.state.dify"],
            ["Task 2.1.4", "test_dify_client.py", "2 用例通过"],
            ["Task 2.2.1", "models/{app,dify_api_key}.py", "2 张表"],
            ["Task 2.2.2", "apps/sync.py", "sync_apps_from_dify"],
            ["Task 2.2.3", "schemas/app.py", "AppOut"],
            ["Task 2.2.4", "apps/router.py（/api/apps/me 占位）", "返回启用 App"],
            ["Task 2.2.5", "test_app_sync.py", "Mock 同步测试"],
            ["Task 2.3.1", "models/app_authorization.py", "复合 PK + CheckConstraint"],
            ["Task 2.3.2", "apps/authorization.py", "grant / revoke / user_can_access_app"],
            ["Task 2.3.3", "test_authorization.py（5 场景）", "全过"],
            ["Task 2.3.4", "/api/apps/me 升级", "按用户解析可见列表"],
        ],
        col_widths=[3.0, 7.0, 5.3],
        first_col_bold=True,
    )

    add_h2(doc, "A.4 阶段 3 任务")
    add_data_table(doc,
        headers=["任务", "步骤", "产出"],
        rows=[
            ["Task 3.1.1", "models/{conversation,message,audit_log}.py", "3 张表"],
            ["Task 3.1.2", "schemas/{conversation,message}.py", "DTO"],
            ["Task 3.1.3", "alembic 迁移（dedupe_key UNIQUE）", "表创建"],
            ["Task 3.2.1", "chat/service.py（user_msg 落库 + 权限校验）", "service 可用"],
            ["Task 3.2.2", "chat/router.py（POST /api/chat/send）", "StreamingResponse"],
            ["Task 3.2.3", "_dify_event_generator", "aiter_lines + finally"],
            ["Task 3.2.4", "超时 + 错误兜底", "event:error + 审计"],
            ["Task 3.2.5", "test_chat_sse.py", "透传 + 落库测试"],
            ["Task 3.3.1", "conversations/router.py", "list / get / 分页"],
            ["Task 3.3.2", "conversations/queries.py（iter_messages）", "流式分批"],
            ["Task 3.3.3", "core/rate_limit.py（内存令牌桶）", "MVP 限流"],
            ["Task 3.4.1", "chat/exports.py（单会话导出）", "GET /api/conversations/{id}/export"],
            ["Task 3.4.2", "conversations/batch_export.py", "POST /api/conversations/export-batch"],
            ["Task 3.4.3", "admin/audit_export.py", "POST /api/admin/audit-logs/export"],
            ["Task 3.4.4", "files/downloads.py", "GET /api/files/{file_id}/download"],
            ["Task 3.4.5", "utils/filename.py", "sanitize_filename"],
            ["Task 3.4.6", "audit/service.py", "audit_log 统一签名"],
            ["Task 3.5.1", "internal/webhook/dify.py", "app-events + workflow-completed"],
            ["Task 3.5.2", "messages dedupe_key 唯一性", "INSERT ON CONFLICT"],
        ],
        col_widths=[3.0, 7.0, 5.3],
        first_col_bold=True,
    )

    add_h2(doc, "A.5 阶段 4 任务")
    add_data_table(doc,
        headers=["任务", "步骤", "产出"],
        rows=[
            ["Task 4.1.1", "files/upload.py（multipart 接收）", "POST /api/chat/files"],
            ["Task 4.1.2", "Content-Length + 流式校验", "20MB 上限"],
            ["Task 4.1.3", "MIME 白名单校验", "pdf/docx/txt/md/png/jpeg"],
            ["Task 4.1.4", "secure_filename 清洗", "路径穿越防御"],
            ["Task 4.1.5", "存本地 + 转 Dify /files/upload", "dify_file_id"],
            ["Task 4.1.6", "上传审计", "audit_logs 落库"],
            ["Task 4.1.7", "test_file_upload.py（4 场景）", "全过"],
        ],
        col_widths=[3.0, 7.0, 5.3],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_appendix_b(doc):
    add_h1(doc, "附录 B  关键路径与甘特图说明")

    add_h2(doc, "B.1 关键路径（Critical Path）")
    add_quote(doc,
        "0.1 → 0.2 → 1.1 → 1.2 → 1.3 → 1.4 → 2.1 → 2.2 → 2.3 → 3.1 → 3.2 → 3.3 → 3.4 → 4.1 → 5.x → 6.x → 7.x\n\n"
        "总工期：10 周")

    add_h2(doc, "B.2 甘特图（文字版）")
    add_code_block(doc, """
       W1   W2   W3   W4   W5   W6   W7   W8   W9   W10
G0:    [==]
G1:        [=====]
G2:                  [==]
G3:                     [==========]
G4:                                   [=]
G5:                                        [========]
G6:                                                   [==]
G7:                                                      [==]

里程碑:  M1  M2    M3    M4        M5  M6       M7  M8
                                (阶段 0-1 衔接)  (阶段 5 末)  (上线)
    """, language="ASCII 甘特图")

    add_h3(doc, "B.2.1 阶段叠加视图")
    add_data_table(doc,
        headers=["周次", "主任务", "并行任务"],
        rows=[
            ["W1", "G0：仓库 + Dify 部署", "—"],
            ["W2", "G1.1 DB + G1.2 JWT", "G0 收尾"],
            ["W3", "G1.3 Auth API + G1.4 组织架构", "G2.1 Dify 客户端（BE-2 兼职）"],
            ["W4", "G2.2 Agent 同步 + G2.3 授权", "G3.1 数据模型"],
            ["W5", "G3.2 SSE", "G3.3 历史查询"],
            ["W6", "G3.4 加密导出后端", "G4.1 文件上传"],
            ["W7", "G5.1 前端骨架 + G5.2 登录", "—"],
            ["W8", "G5.3 对话 UI + G5.4 加密导出前端", "—"],
            ["W9", "G6 管理员后台", "—"],
            ["W10", "G7 部署 + 验收", "—"],
        ],
        col_widths=[1.5, 6.5, 7.3],
        first_col_bold=True,
    )

    add_h3(doc, "B.2.2 资源负载图（人周）")
    add_data_table(doc,
        headers=["角色", "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8", "W9", "W10", "合计"],
        rows=[
            ["PM", "0.5", "0.5", "0.5", "0.5", "0.5", "0.5", "0.5", "0.5", "0.5", "1.0", "5.5"],
            ["BE-1", "1.0", "1.0", "1.0", "1.0", "1.0", "1.0", "0.5", "0.5", "1.0", "0.5", "8.5"],
            ["BE-2", "0.3", "0.5", "1.0", "1.0", "1.0", "1.0", "0.5", "0.5", "0.5", "0.5", "6.8"],
            ["FE-1", "0.3", "0.3", "0.3", "0.3", "0.3", "0.3", "1.0", "1.0", "1.0", "0.5", "5.3"],
            ["SRE", "1.0", "0.1", "0.1", "0.1", "0.1", "0.1", "0.1", "0.1", "0.1", "1.0", "2.8"],
            ["安全", "—", "0.3", "—", "0.3", "0.3", "0.3", "—", "—", "—", "0.5", "1.7"],
            ["测试", "—", "—", "—", "—", "0.5", "0.5", "0.5", "0.5", "0.5", "1.0", "3.5"],
            ["业务", "0.3", "0.3", "0.3", "0.3", "0.3", "0.3", "0.3", "0.3", "0.3", "0.5", "3.2"],
            ["合计", "3.4", "3.0", "3.2", "3.5", "4.0", "4.0", "3.4", "3.4", "3.9", "5.5", "37.3"],
        ],
        col_widths=[1.5] + [0.9] * 11,
        first_col_bold=True,
    )

    add_quote(doc,
        "上表为人周投入（兼职按比例计算），实际工时（人天）= 人周 × 5。\n"
        "峰值：W10（5.5 人周）— 部署 + 验收 + 安全 review 同时进行。")

    add_page_break(doc)


def write_appendix_c(doc):
    add_h1(doc, "附录 C  文档索引")

    add_h2(doc, "C.1 项目文档清单")
    add_data_table(doc,
        headers=["#", "文档名", "路径", "用途"],
        rows=[
            ["D-01", "企业内部 Agent 平台设计文档", "docs/superpowers/specs/2026-08-28-agent-platform-design.md",
             "总体技术设计"],
            ["D-02", "MVP Phase 1 实施计划", "docs/superpowers/plans/2026-08-28-mvp-phase-1-infra-auth.md",
             "后端任务详细步骤"],
            ["D-03", "实施环境准备文档", "docs/superpowers/plans/2026-08-28-environment-setup.md",
             "环境搭建操作步骤"],
            ["D-04", "加密导出特性设计", "docs/superpowers/specs/2026-08-28-encrypted-export-design.md",
             "加密导出设计"],
            ["D-05", "建设技术方案框架", "output/Agent平台建设技术方案框架.docx",
             "总体技术方案"],
            ["D-06", "项目执行计划（本文件）", "output/Agent平台执行计划.docx",
             "执行计划（当前）"],
        ],
        col_widths=[1.0, 4.5, 6.5, 3.3],
        first_col_bold=True,
    )

    add_h2(doc, "C.2 文档版本控制")
    add_data_table(doc,
        headers=["版本", "日期", "变更说明", "作者"],
        rows=[
            ["V0.1", "2026-08-25", "初稿，章节结构草案", "PM"],
            ["V0.5", "2026-08-27", "完成阶段划分 + 任务分解", "PM + BE-1"],
            ["V0.9", "2026-08-28", "完成风险登记册 + 里程碑", "PM + 安全"],
            ["V1.0", "2026-08-28", "正式发布（待评审）", "PM"],
        ],
        col_widths=[2.0, 2.5, 7.0, 3.8],
        first_col_bold=True,
    )

    add_h2(doc, "C.3 文档评审与签发")
    add_data_table(doc,
        headers=["评审项", "评审人", "状态"],
        rows=[
            ["技术方案可行性", "应用开发组组长", "待评审"],
            ["进度合理性", "数字化办公室主任", "待评审"],
            ["资源 / 预算", "基础架构组组长", "待评审"],
            ["安全风险评估", "安全团队代表", "待评审"],
            ["业务需求对齐", "业务代表", "待评审"],
            ["变更管理流程", "CCB 主席", "待评审"],
        ],
        col_widths=[5.0, 5.0, 5.3],
        first_col_bold=True,
    )

    # 文档结束
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=24, after=0, line_spacing=1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 全文完 ——")
    set_run_font(run, size_pt=12, bold=True, color="888888")


# ============================================================================
# 主入口
# ============================================================================

def main():
    output_dir = Path("/mnt/e/program/agent_platform/output")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "Agent平台执行计划.docx"

    doc = init_document()

    # 自定义页眉文字（覆盖默认）
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    # 清空原有 run
    for run in list(header_para.runs):
        run.text = ""
    header_run = header_para.add_run("企业内部 Agent 智能体平台  项目执行计划")
    set_run_font(header_run, size_pt=9, color="888888")

    # 封面
    add_cover(doc)

    # 目录
    add_toc(doc)

    # 章节
    write_chapter_1(doc)
    write_chapter_2(doc)
    write_chapter_3(doc)
    write_chapter_4(doc)
    write_chapter_5(doc)
    write_chapter_6(doc)
    write_chapter_7(doc)
    write_chapter_8(doc)
    write_chapter_9(doc)
    write_chapter_10(doc)
    write_chapter_11(doc)
    write_appendix_a(doc)
    write_appendix_b(doc)
    write_appendix_c(doc)

    doc.save(str(output_path))
    print(f"✅ 执行计划 Word 文档已生成：{output_path}")
    print(f"   文件大小：{output_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
