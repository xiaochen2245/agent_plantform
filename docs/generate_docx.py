"""生成企业内部 Agent 智能体平台建设技术方案框架 Word 文档。

Usage:
    python generate_docx.py

Output:
    /mnt/e/program/agent_platform/output/Agent平台建设技术方案框架.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============================================================================
# 1. 样式与字体设置
# ============================================================================

def set_cell_bg(cell, color_hex: str):
    """设置单元格背景色（浅灰）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, name_cn: str = "微软雅黑", name_en: str = "Calibri",
                 size_pt: float = 11, bold: bool = False, color: str | None = None):
    """设置 run 的中英文字体、字号、加粗、颜色。"""
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    rFonts.set(qn('w:cs'), name_en)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6,
                          line_spacing: float = 1.5):
    """设置段前段后间距 + 行距。"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_page_break(doc):
    """添加分页符。"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    # 用 WD_BREAK.PAGE 替代
    from docx.enum.text import WD_BREAK
    p.runs[0].add_break(WD_BREAK.PAGE)


def add_horizontal_line(paragraph):
    """为段落添加底部边框线（用于分隔）。"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_table_borders(table, color: str = "BFBFBF", size: str = "4"):
    """为表格设置边框。"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ============================================================================
# 2. 文档初始化 & 样式预设
# ============================================================================

def init_document() -> Document:
    """初始化文档：A4、页边距、页眉页脚。"""
    doc = Document()

    # A4 页面 + 适中页边距
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 默认正文字体（Normal 样式）
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')

    # 页脚：页码
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    set_run_font(run, size_pt=9, color="666666")

    # 插入 PAGE 域
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE   \\* MERGEFORMAT'
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)

    # 页眉：左侧文档标题
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run("企业内部 Agent 智能体平台  建设技术方案框架")
    set_run_font(header_run, size_pt=9, color="888888")

    return doc


# ============================================================================
# 3. 标题 & 正文辅助函数
# ============================================================================

def add_h1(doc, text: str):
    """一级标题（章）：黑体 18pt，居中，段前 18pt，段后 12pt。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=18, after=12, line_spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, name_cn="黑体", size_pt=18, bold=True, color="1F4E79")
    # 添加 OutlineLevel 0（用于导航/目录）
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), '0')
    pPr.append(outline)
    # 段后加线
    add_horizontal_line(p)
    return p


def add_h2(doc, text: str):
    """二级标题（节）：黑体 14pt，左对齐，段前 12pt，段后 6pt。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=6, line_spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, name_cn="黑体", size_pt=14, bold=True, color="2E74B5")
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), '1')
    pPr.append(outline)
    return p


def add_h3(doc, text: str):
    """三级标题（小节）：黑体 12pt，段前 8pt，段后 4pt。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line_spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, name_cn="黑体", size_pt=12, bold=True, color="404040")
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), '2')
    pPr.append(outline)
    return p


def add_h4(doc, text: str):
    """四级标题：粗体 11pt，段前 4pt，段后 2pt。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=2, line_spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, size_pt=11, bold=True, color="404040")
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), '3')
    pPr.append(outline)
    return p


def add_body(doc, text: str, indent_first: bool = True, bold: bool = False):
    """正文段落：1.5 倍行距，首行缩进 2 字符。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line_spacing=1.5)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 约 2 个字符
    run = p.add_run(text)
    set_run_font(run, size_pt=11, bold=bold)
    return p


def add_bullet(doc, text: str, level: int = 0):
    """项目符号列表。"""
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.4)
    p.paragraph_format.left_indent = Cm(0.74 + 0.6 * level)
    run = p.add_run(text)
    set_run_font(run, size_pt=11)
    return p


def add_number(doc, text: str):
    """编号列表。"""
    p = doc.add_paragraph(style='List Number')
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.4)
    run = p.add_run(text)
    set_run_font(run, size_pt=11)
    return p


def add_code_block(doc, code: str, language: str = ""):
    """代码块：等宽字体 + 浅灰背景 + 段前段后间距。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line_spacing=1.15)
    if language:
        # 语言标签单独一段
        lang_p = doc.add_paragraph()
        lang_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lang_run = lang_p.add_run(f"【{language}】")
        set_run_font(lang_run, size_pt=10, bold=True, color="666666")
        set_paragraph_spacing(lang_p, before=4, after=2, line_spacing=1.0)
    run = p.add_run(code)
    set_run_font(run, name_cn="Consolas", name_en="Consolas", size_pt=9, color="2D2D2D")
    # 浅灰背景 + 缩进
    pPr = p._p.get_or_add_pPr()
    # 段落底纹
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    # 段落缩进
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '300')
    ind.set(qn('w:right'), '300')
    pPr.append(ind)
    # 段落边框
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'BFBFBF')
        pBdr.append(border)
    pPr.append(pBdr)
    return p


def add_quote(doc, text: str):
    """引用块：缩进 + 斜体 + 左侧蓝色边。"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=4, line_spacing=1.4)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.right_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size_pt=10.5, color="404040")
    # 左侧边框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '2E74B5')
    pBdr.append(left)
    pPr.append(pBdr)
    return p


# ============================================================================
# 4. 表格辅助函数
# ============================================================================

def add_data_table(doc, headers: list[str], rows: list[list[str]],
                   header_bg: str = "1F4E79", first_col_bold: bool = False,
                   col_widths: list[float] | None = None):
    """创建带表头的数据表格。

    Args:
        headers: 表头列表
        rows: 行数据列表
        header_bg: 表头背景色（深蓝）
        first_col_bold: 首列是否加粗
        col_widths: 各列宽度（cm），None 则均分
    """
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    set_table_borders(table, color="808080", size="6")

    # 列宽设置
    if col_widths is None:
        col_widths = [16.0 / len(headers)] * len(headers)
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, name_cn="黑体", size_pt=10.5, bold=True, color="FFFFFF")
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            bold = first_col_bold and c_idx == 0
            set_run_font(run, size_pt=10, bold=bold)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 斑马纹
            if (r_idx + 1) % 2 == 0:
                set_cell_bg(cell, "F2F7FB")

    # 表格后段距
    p_after = doc.add_paragraph()
    set_paragraph_spacing(p_after, before=0, after=6, line_spacing=1.0)

    return table


# ============================================================================
# 5. 封面
# ============================================================================

def add_cover(doc: Document):
    """添加封面。"""
    # 顶部留白
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=24, line_spacing=1.5)
    run = p.add_run("企业内部 Agent 智能体平台")
    set_run_font(run, name_cn="黑体", size_pt=32, bold=True, color="1F4E79")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=48, line_spacing=1.5)
    run = p.add_run("建 设 技 术 方 案 框 架")
    set_run_font(run, name_cn="黑体", size_pt=22, bold=True, color="2E74B5")

    # 分隔线
    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sep_p, before=0, after=24, line_spacing=1.0)
    pPr = sep_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=24, after=12, line_spacing=1.5)
    run = p.add_run("—— 基于 Dify 社区版的智能对话门户建设方案 ——")
    set_run_font(run, size_pt=14, color="404040")

    # 元信息表格
    for _ in range(2):
        doc.add_paragraph()

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.columns[0].width = Cm(4)
    meta_table.columns[1].width = Cm(8)
    meta_data = [
        ("文档版本", "V1.0"),
        ("编写日期", "2026-08-28"),
        ("文档状态", "待评审"),
        ("密级", "内部公开"),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c1 = row.cells[0]
        c2 = row.cells[1]
        c1.width = Cm(4)
        c2.width = Cm(8)
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(k)
        set_run_font(run1, name_cn="黑体", size_pt=12, bold=True, color="FFFFFF")
        set_cell_bg(c1, "1F4E79")
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(v)
        set_run_font(run2, size_pt=12, color="1F4E79")
        set_cell_bg(c2, "F2F7FB")
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(meta_table, color="1F4E79", size="8")

    # 底部组织名
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=4, line_spacing=1.5)
    run = p.add_run("企业数字化转型办公室")
    set_run_font(run, size_pt=14, bold=True, color="404040")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)
    run = p.add_run("2026 年 8 月")
    set_run_font(run, size_pt=12, color="666666")

    add_page_break(doc)


# ============================================================================
# 6. 目录（手写）
# ============================================================================

def add_toc(doc: Document):
    """手写目录占位（Word 中可手动更新）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=18, line_spacing=1.5)
    run = p.add_run("目  录")
    set_run_font(run, name_cn="黑体", size_pt=20, bold=True, color="1F4E79")
    add_horizontal_line(p)

    toc_entries = [
        ("第一章  项目概述", "1"),
        ("  1.1 项目背景", "1"),
        ("  1.2 建设目标", "1"),
        ("  1.3 建设范围", "2"),
        ("  1.4 关键决策", "2"),
        ("第二章  需求分析", "3"),
        ("  2.1 业务需求", "3"),
        ("  2.2 功能需求", "3"),
        ("  2.3 非功能需求", "4"),
        ("第三章  总体架构设计", "5"),
        ("  3.1 架构总览", "5"),
        ("  3.2 分层架构说明", "5"),
        ("  3.3 技术栈选型", "6"),
        ("第四章  关键技术方案", "8"),
        ("  4.1 鉴权与权限模型", "8"),
        ("  4.2 Dify 集成方案", "10"),
        ("  4.3 对话流与 SSE 代理", "12"),
        ("  4.4 数据模型设计", "15"),
        ("  4.5 文件上传方案", "17"),
        ("  4.6 加密导出方案", "18"),
        ("第五章  部署架构", "22"),
        ("  5.1 网络拓扑", "22"),
        ("  5.2 Docker 编排", "22"),
        ("  5.3 启动顺序", "23"),
        ("  5.4 关键配置", "24"),
        ("第六章  安全方案", "26"),
        ("  6.1 身份认证安全", "26"),
        ("  6.2 CSRF 防护", "26"),
        ("  6.3 文件上传安全", "27"),
        ("  6.4 数据加密", "28"),
        ("  6.5 审计日志", "28"),
        ("第七章  测试方案", "30"),
        ("  7.1 测试策略", "30"),
        ("  7.2 单元测试", "30"),
        ("  7.3 集成测试", "31"),
        ("  7.4 E2E 测试", "31"),
        ("第八章  项目实施计划", "33"),
        ("  8.1 阶段划分", "33"),
        ("  8.2 资源投入", "34"),
        ("  8.3 里程碑", "34"),
        ("第九章  风险与对策", "35"),
        ("  9.1 技术风险", "35"),
        ("  9.2 进度风险", "35"),
        ("  9.3 安全风险", "36"),
        ("附录 A  术语表", "37"),
        ("附录 B  参考文献", "37"),
    ]
    for title, page in toc_entries:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line_spacing=1.4)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)  # leader=2 → 点线
        # 标题
        run = p.add_run(title)
        if title.startswith("第") and "章" in title[:3]:
            set_run_font(run, name_cn="黑体", size_pt=11, bold=True)
        else:
            set_run_font(run, size_pt=10.5)
        # 制表符
        tab_run = p.add_run("\t")
        # 页码
        page_run = p.add_run(page)
        set_run_font(page_run, size_pt=10.5)

    add_page_break(doc)


# ============================================================================
# 7. 各章节内容
# ============================================================================

def write_chapter_1(doc):
    add_h1(doc, "第一章  项目概述")

    add_h2(doc, "1.1 项目背景")
    add_body(doc,
        "随着大语言模型（LLM）技术的快速发展和企业内部知识管理需求的日益增长，"
        "传统搜索引擎和知识库已难以满足员工对\"精准问答 + 业务推理 + 知识检索\"的复合诉求。"
        "近两年来，以 GPT、Claude 等为代表的大模型在企业内部多个场景展现出巨大潜力，"
        "但同时也面临模型选型困难、知识更新滞后、合规风险高、权限管控复杂等挑战。")

    add_body(doc,
        "当前公司在文档检索、IT 工单、HR 答疑、技术问答等高频场景中，员工日均产生大量重复性问题，"
        "严重占用业务专家时间。现有方案存在三大痛点：")

    add_number(doc, "分散的 AI 入口：各部门各自对接外部或私有化 LLM，缺乏统一管控；")
    add_number(doc, "缺乏权限隔离：业务敏感信息无法按部门 / 角色做精细化授权；")
    add_number(doc, "审计与合规困难：员工与 AI 的对话内容缺乏结构化记录，难以满足合规报送需求。")

    add_body(doc,
        "本项目旨在构建一个企业内部的 Agent 智能体平台，统一对外提供 AI 对话入口，"
        "并通过自研的鉴权 + 代理 + 审计层，将底层 LLM 引擎（Dify 社区版）的能力"
        "以安全、可控、可审计的方式开放给公司全体员工。")

    add_h2(doc, "1.2 建设目标")
    add_body(doc, "本项目总体目标是：构建一个面向公司全体员工的统一 AI 对话门户，"
                  "支持多 Agent（智能体）授权、SSE 流式对话、文件安全上传、加密导出、完整审计等核心能力，"
                  "为各业务部门提供\"开箱即用、按需授权、可管可控\"的 AI 服务底座。")

    add_h3(doc, "1.2.1 总体目标")
    add_data_table(doc,
        headers=["目标维度", "具体目标", "衡量指标"],
        rows=[
            ["统一入口", "面向员工提供统一的 Web 门户", "Portal 上线，员工登录可达"],
            ["多 Agent 复用", "支持对接 5+ 业务 Agent（HR、IT、研发知识库等）", "管理员可自助授权"],
            ["权限隔离", "按用户 / 部门 / 角色三态授权", "权限校验覆盖率达 100%"],
            ["对话体验", "首字节延迟 ≤ 300ms", "P95 首字节延迟 ≤ 500ms"],
            ["审计合规", "所有敏感操作有日志，可导出", "审计覆盖关键路径 100%"],
            ["文件安全", "支持常见办公文档 / 图片上传", "大小 / MIME / 文件名三重校验"],
            ["数据自主", "会话可加密导出 / 解密", "密码不出客户端"],
        ],
        col_widths=[2.8, 7.0, 4.5],
        first_col_bold=True,
    )

    add_h3(doc, "1.2.2 阶段目标")
    add_bullet(doc, "MVP（6–10 周）：完成 Dify 集成、JWT 鉴权、多 Agent 授权、SSE 对话、文件上传、基础审计、加密导出；")
    add_bullet(doc, "二期（视需求）：限流、用量配额、用户反馈、统计视图、SSE 断点续传、ClamAV 病毒扫描；")
    add_bullet(doc, "三期（远期）：多模态上传预览、SSO 接入、跨对话知识检索、外部知识源集成。")

    add_h2(doc, "1.3 建设范围")
    add_body(doc, "本项目 MVP 阶段覆盖以下范围：")

    add_h3(doc, "1.3.1 功能范围")
    add_data_table(doc,
        headers=["功能模块", "范围", "交付物"],
        rows=[
            ["门户前端", "登录 / 对话 / 历史 / 管理员后台（用户管理 + Agent 授权）", "React SPA + Ant Design"],
            ["后端服务", "鉴权 + 权限模型 + Agent 同步 + 对话代理 + SSE + 审计 + 文件上传 + 加密导出", "FastAPI + SQLAlchemy"],
            ["Dify 引擎", "复用社区版作为 Agent 编排 + RAG 引擎（独立部署）", "Docker Compose 部署"],
            ["基础设施", "Postgres 16 + Nginx 反向代理 + Docker 编排", "deploy/ 目录"],
        ],
        col_widths=[3.5, 7.5, 4.5],
        first_col_bold=True,
    )

    add_h3(doc, "1.3.2 不在本期范围")
    add_number(doc, "企业 SSO / LDAP 接入（接口预留，代码层级未实现）")
    add_number(doc, "移动端 App（仅 Web 端 MVP）")
    add_number(doc, "PDF / Word 格式导出（仅 JSON / Markdown）")
    add_number(doc, "跨对话知识检索（pgvector 二期）")
    add_number(doc, "用量配额与限流（Redis 二期）")

    add_h2(doc, "1.4 关键决策")
    add_quote(doc,
        "本节列出的关键决策均经过多轮评审，是后续设计与实施的指导原则。"
        "后续如有变更需走正式评审流程，避免实施阶段反复返工。")

    add_data_table(doc,
        headers=["决策项", "决策", "理由"],
        rows=[
            ["底层引擎", "复用 Dify 社区版作为 Agent 编排 + RAG 引擎",
             "社区版成熟度高、自带工作流画布与知识库；自研 Agent 引擎投入产出比不合理"],
            ["后端框架", "FastAPI（Python 3.11+）",
             "异步生态完善、类型提示友好、SSE 支持原生、社区活跃"],
            ["前端框架", "React 18 + TypeScript + Vite + Ant Design 5",
             "公司技术栈匹配、生态成熟、AI 场景社区组件丰富"],
            ["数据库", "PostgreSQL 16",
             "JSONB / 物化路径 / 全文检索能力齐全，与 FastAPI 异步适配良好"],
            ["鉴权方案", "JWT（access 15min + refresh 7d）+ httpOnly + SameSite=Strict cookie",
             "无状态、易扩展；refresh token 入库支持主动撤销"],
            ["部署形态", "Docker Compose 一把梭（Postgres / Dify / Nginx / 后端 / 前端）",
             "环境一致性高、运维成本低、便于内网部署"],
            ["加密导出", "浏览器侧 Web Crypto API + AES-256-GCM + PBKDF2 100k 轮",
             "密码不出客户端、支持离线解密 /decrypt 页面"],
            ["MCP 服务预留", "接口预留 /api/auth/oauth/{provider}/callback",
             "二期接企业 SSO 时不破坏现有结构"],
        ],
        col_widths=[3.0, 5.5, 8.0],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_chapter_2(doc):
    add_h1(doc, "第二章  需求分析")

    add_h2(doc, "2.1 业务需求")
    add_body(doc,
        "业务层面，本项目需解决三类核心诉求：")

    add_h3(doc, "2.1.1 统一 AI 入口")
    add_body(doc,
        "为员工提供单一 Web 入口访问 AI 服务，所有 Agent（HR 问答、IT 工单、研发知识库、"
        "运营答疑等）通过统一门户暴露，员工无需关心底层 LLM 实现，"
        "管理员可通过后台灵活开通 / 关闭 Agent。")

    add_h3(doc, "2.1.2 精细化权限管控")
    add_body(doc,
        "支持按用户、部门、角色三种维度授权，满足\"某 Agent 仅限财务部 + 财务总监可见\""
        "等典型合规诉求。授权关系变更须支持即时生效，避免策略生效延迟导致越权访问。")

    add_h3(doc, "2.1.3 可审计可追溯")
    add_body(doc,
        "所有敏感操作（登录、对话发起、文件上传、Agent 授权变更、对话导出等）"
        "需落入审计日志，支持按时间窗口导出，用于合规报送与事后追查。"
        "审计日志不存对话原文，仅存元数据，避免敏感信息二次泄露。")

    add_h2(doc, "2.2 功能需求")
    add_body(doc, "本节按模块梳理 MVP 阶段的功能需求清单。")

    add_h3(doc, "2.2.1 鉴权与权限模块")
    add_data_table(doc,
        headers=["需求编号", "功能描述", "优先级"],
        rows=[
            ["FR-AUTH-01", "邮箱 + 密码登录（bcrypt 哈希）", "P0"],
            ["FR-AUTH-02", "JWT 颁发（access 15min + refresh 7d）", "P0"],
            ["FR-AUTH-03", "Refresh token 入库 + 主动撤销 + 轮转", "P0"],
            ["FR-AUTH-04", "CSRF 防护（SameSite=Strict + Origin 校验）", "P0"],
            ["FR-AUTH-05", "JWT 载荷含 user_id / roles[] / dept_id / jti", "P0"],
            ["FR-AUTH-06", "OAuth 接口预留（/api/auth/oauth/{provider}/callback）", "P1"],
            ["FR-AUTH-07", "Token 黑名单 / 主动下线（二期）", "P2"],
        ],
        col_widths=[2.5, 11.0, 1.8],
        first_col_bold=True,
    )

    add_h3(doc, "2.2.2 对话模块")
    add_data_table(doc,
        headers=["需求编号", "功能描述", "优先级"],
        rows=[
            ["FR-CHAT-01", "SSE 流式对话（首字节延迟 ≤ 300ms）", "P0"],
            ["FR-CHAT-02", "用户可见 Agent 列表按权限过滤", "P0"],
            ["FR-CHAT-03", "FastAPI 侧二次权限校验（防绕过前端）", "P0"],
            ["FR-CHAT-04", "多轮对话（Dify conversation_id 关联）", "P0"],
            ["FR-CHAT-05", "用户消息 + Agent 回复落库", "P0"],
            ["FR-CHAT-06", "流中断兜底（finally 块强制写库）", "P0"],
            ["FR-CHAT-07", "SSE 断点续传（二期）", "P2"],
        ],
        col_widths=[2.5, 11.0, 1.8],
        first_col_bold=True,
    )

    add_h3(doc, "2.2.3 文件上传模块")
    add_data_table(doc,
        headers=["需求编号", "功能描述", "优先级"],
        rows=[
            ["FR-FILE-01", "大小上限 20MB（Content-Length + 流式双重校验）", "P0"],
            ["FR-FILE-02", "MIME 白名单（pdf/docx/txt/md/png/jpeg）", "P0"],
            ["FR-FILE-03", "文件名清洗（防路径穿越 / CRLF 注入）", "P0"],
            ["FR-FILE-04", "上传审计（user_id / file_name / size / mime / sha256）", "P0"],
            ["FR-FILE-05", "转发到 Dify /files/upload 拿 dify_file_id", "P0"],
            ["FR-FILE-06", "ClamAV 病毒扫描（二期）", "P2"],
        ],
        col_widths=[2.5, 11.0, 1.8],
        first_col_bold=True,
    )

    add_h3(doc, "2.2.4 加密导出模块")
    add_data_table(doc,
        headers=["需求编号", "功能描述", "优先级"],
        rows=[
            ["FR-EXP-01", "单会话导出（JSON / Markdown）", "P0"],
            ["FR-EXP-02", "批量会话导出（zip，≤ 100 个）", "P0"],
            ["FR-EXP-03", "审计日志导出（CSV / JSON）", "P0"],
            ["FR-EXP-04", "附件下载（带 sanitize_filename）", "P0"],
            ["FR-EXP-05", "浏览器侧 AES-256-GCM + PBKDF2 100k 加密", "P0"],
            ["FR-EXP-06", "公共 /decrypt 页面（envelope 校验 + 离线解密）", "P0"],
            ["FR-EXP-07", "密码强度评估（zxcvbn） + 弱密码警告", "P0"],
        ],
        col_widths=[2.5, 11.0, 1.8],
        first_col_bold=True,
    )

    add_h2(doc, "2.3 非功能需求")
    add_body(doc, "除功能需求外，本项目还需满足以下非功能需求。")

    add_h3(doc, "2.3.1 性能需求")
    add_data_table(doc,
        headers=["指标", "目标", "说明"],
        rows=[
            ["首字节延迟", "P50 ≤ 300ms，P95 ≤ 500ms", "SSE 流式对话"],
            ["页面加载", "P95 ≤ 2s", "门户前端首屏"],
            ["并发对话", "≥ 200 路", "FastAPI 2 实例 × uvicorn 2 workers"],
            ["审计查询", "≤ 1s", "单次按时间窗口查询"],
        ],
        col_widths=[4.0, 5.0, 6.3],
        first_col_bold=True,
    )

    add_h3(doc, "2.3.2 可用性需求")
    add_bullet(doc, "FastAPI 2 容器实例 + Nginx 负载均衡，单实例故障不影响整体可用；")
    add_bullet(doc, "关键路径（登录、对话、审计）保证 99.5% 可用；")
    add_bullet(doc, "Postgres 每日自动备份，保留 7 天；")
    add_bullet(doc, "Dify / 后端 / 前端均支持滚动升级。")

    add_h3(doc, "2.3.3 安全需求")
    add_bullet(doc, "所有敏感操作（写操作）必须有身份验证 + CSRF 防护；")
    add_bullet(doc, "密码 bcrypt 哈希，禁止明文存储；")
    add_bullet(doc, "JWT 短期（15min），refresh token 主动撤销；")
    add_bullet(doc, "Dify API Key Fernet 加密存储，与 JWT 密钥独立；")
    add_bullet(doc, "文件上传大小 / MIME / 文件名三重校验；")
    add_bullet(doc, "对话导出可加密，密码不出客户端。")

    add_h3(doc, "2.3.4 可维护性需求")
    add_bullet(doc, "日志结构化（JSON）+ 关键字段（worker_id / user_id / conv_id / app_id / duration_ms）；")
    add_bullet(doc, "所有环境变量通过 .env 管理，密钥不进 git；")
    add_bullet(doc, "Docker 镜像版本锁定（次版本号 + 关键镜像 digest）；")
    add_bullet(doc, "数据库迁移通过 Alembic 管理，可前滚 / 回退。")

    add_h3(doc, "2.3.5 合规需求")
    add_bullet(doc, "审计日志保留 ≥ 180 天（满足一般合规要求）；")
    add_bullet(doc, "员工离职触发对话软删 + Dify 账号禁用（接口预留）；")
    add_bullet(doc, "审计导出 CSV 注入防护（公式前缀处理）。")

    add_page_break(doc)


def write_chapter_3(doc):
    add_h1(doc, "第三章  总体架构设计")

    add_h2(doc, "3.1 架构总览")
    add_body(doc,
        "本平台采用三层架构设计：前端（React SPA）作为员工入口，"
        "FastAPI 后端作为鉴权 / 授权 / 代理 / 审计网关，"
        "Dify 社区版作为 Agent 编排与 RAG 引擎。整体架构图如下：")

    add_code_block(doc, """
┌──────────────────────────────────────────────────────────┐
│ 浏览器 (React SPA)                                       │
│   - Ant Design UI                                        │
│   - axios + fetch + ReadableStream (SSE 消费)            │
│   - JWT 存 httpOnly cookie                               │
└──────────────────────────┬───────────────────────────────┘
                         │ HTTPS / SSE
                         ▼
┌──────────────────────────────────────────────────────────┐
│  FastAPI 后端（BFF + 业务层）                            │
│   - /api/auth/*          JWT 登录、刷新、注销           │
│   - /api/apps/*          用户可见 App 列表、权限过滤    │
│   - /api/chat/*          对话消息、流式响应             │
│   - /api/conversations/* 对话历史（从本地 DB 查）       │
│   - /api/admin/*         用户 / 部门 / 角色管理         │
│   - /internal/dify/*     Dify 适配层（API Key 注入、转发）│
│   - /internal/webhook/*  Dify 事件回调                  │
└──────────────┬──────────────────────┬────────────────────┘
               │                      │
               ▼                      ▼
   ┌───────────────────┐   ┌───────────────────────────────┐
   │ Postgres 16       │   │ Dify 社区版（独立部署）        │
   │ - users           │   │ - 工作流编排画布              │
   │ - departments     │   │ - RAG / 知识库管理            │
   │ - roles           │   │ - 应用 API (chat-messages)    │
   │ - user_roles      │   │ - Webhook → 我们回调          │
   │ - apps            │   └───────────────────────────────┘
   │ - app_authorizations                              │
   │ - dify_api_keys                                   │
   │ - conversations                                   │
   │ - messages                                        │
   │ - audit_logs                                      │
   └───────────────────┘
    """, language="架构示意图")

    add_h2(doc, "3.2 分层架构说明")
    add_body(doc, "本节按职责对三层架构中的每一层做详细说明。")

    add_h3(doc, "3.2.1 前端层（员工入口）")
    add_body(doc, "前端为 React 单页应用，承担以下职责：")
    add_bullet(doc, "员工身份登录 / 注销，JWT 由 httpOnly cookie 自动管理；")
    add_bullet(doc, "按权限显示可访问的 Agent 列表；")
    add_bullet(doc, "对话流式 UI（fetch + ReadableStream，不使用 EventSource 以避免断线重发整条消息）；")
    add_bullet(doc, "对话历史查询（从 FastAPI 查本地 DB）；")
    add_bullet(doc, "管理员后台（用户 / 部门 / 角色 / Agent 授权）；")
    add_bullet(doc, "加密下载 UI（PasswordDialog + DownloadButton + useEncryptedDownload hook）。")

    add_h3(doc, "3.2.2 后端层（鉴权 + 代理 + 审计网关）")
    add_body(doc, "FastAPI 后端承担平台核心业务逻辑，是整个系统的安全屏障：")
    add_bullet(doc, "JWT 登录 / 刷新 / 注销；")
    add_bullet(doc, "权限解析（用户授权 ∪ 部门授权 ∪ 角色授权）；")
    add_bullet(doc, "Dify API Key 管理（Fernet 加密存储 + 启动解密 + 调用注入）；")
    add_bullet(doc, "SSE 流式代理（Dify → FastAPI → 前端逐 chunk 透传）；")
    add_bullet(doc, "对话消息落库（用户 + Assistant 双写 + dedupe_key 去重）；")
    add_bullet(doc, "文件上传校验 + 转发；")
    add_bullet(doc, "审计日志写入；")
    add_bullet(doc, "加密导出（流式 JSON / Markdown / zip / CSV）。")

    add_h3(doc, "3.2.3 引擎层（Dify 社区版）")
    add_body(doc, "Dify 社区版作为编排与 RAG 引擎，仅供 5–20 名管理员 / 编辑者使用：")
    add_bullet(doc, "工作流编排（chat / completion / workflow / agent 四种模式）；")
    add_bullet(doc, "知识库管理（文档上传、向量化、检索）；")
    add_bullet(doc, "应用 API（chat-messages 流式 / completion 阻塞）；")
    add_bullet(doc, "Webhook 事件（应用 CRUD + 工作流完成 → 兜底补齐 messages）。")

    add_h3(doc, "3.2.4 数据层（Postgres 16）")
    add_body(doc, "Postgres 作为主存储，承载所有业务数据：")
    add_bullet(doc, "用户 / 部门 / 角色 / 多对多关系；")
    add_bullet(doc, "Agent 镜像 + 三态授权关系；")
    add_bullet(doc, "Dify API Key 加密存储；")
    add_bullet(doc, "对话 + 消息双写（UUID 主键）；")
    add_bullet(doc, "审计日志（metadata JSONB）。")

    add_h2(doc, "3.3 技术栈选型")
    add_body(doc, "本节梳理各层技术选型与版本约束，所有版本均锁定到次版本号，"
                  "避免升级引入的隐性破坏。")

    add_h3(doc, "3.3.1 前端技术栈")
    add_data_table(doc,
        headers=["技术", "版本", "用途"],
        rows=[
            ["React", "18.2+", "UI 框架"],
            ["TypeScript", "5.4+", "静态类型"],
            ["Vite", "5.1+", "构建工具"],
            ["Ant Design", "5.15+", "组件库"],
            ["React Router", "6.22+", "路由"],
            ["Zustand", "4.5+", "轻量状态管理"],
            ["Axios", "1.6+", "普通 HTTP 请求"],
            ["原生 fetch + ReadableStream", "浏览器原生", "SSE 流式消费"],
            ["dayjs", "1.11+", "日期处理"],
            ["zxcvbn", "4.4.2", "密码强度评估"],
            ["@transcend-io/penumbra", "8.1.4", "浏览器侧 AES-256-GCM 流式加密"],
        ],
        col_widths=[5.0, 3.0, 7.3],
        first_col_bold=True,
    )

    add_h3(doc, "3.3.2 后端技术栈")
    add_data_table(doc,
        headers=["技术", "版本", "用途"],
        rows=[
            ["Python", "3.11+", "开发语言"],
            ["FastAPI", "0.110+", "Web 框架"],
            ["uvicorn", "0.27+", "ASGI 服务器（2 workers / 实例）"],
            ["SQLAlchemy", "2.0.27+", "ORM（异步）"],
            ["asyncpg", "0.29+", "Postgres 异步驱动"],
            ["Alembic", "1.13+", "数据库迁移"],
            ["Pydantic", "2.6+", "数据校验 + Settings"],
            ["python-jose", "3.3+", "JWT 编解码"],
            ["passlib[bcrypt]", "1.7.4+", "密码哈希"],
            ["httpx", "0.27+", "异步 HTTP 客户端（调 Dify）"],
            ["python-multipart", "0.0.9+", "multipart 解析（文件上传）"],
            ["cryptography", "42.0+", "Fernet 对称加密"],
            ["sse-starlette", "2.0+", "SSE 响应支持"],
            ["loguru", "0.7+", "结构化日志"],
            ["zipstream-ng", "1.1.0", "流式 zip 打包（批量导出）"],
        ],
        col_widths=[5.0, 3.0, 7.3],
        first_col_bold=True,
    )

    add_h3(doc, "3.3.3 基础设施技术栈")
    add_data_table(doc,
        headers=["技术", "版本", "用途"],
        rows=[
            ["Postgres", "16-alpine", "主数据库"],
            ["Redis", "7-alpine", "仅 Dify 使用；MVP 不引入业务限流"],
            ["Nginx", "1.25-alpine", "HTTPS 反代 + 前端静态服务 + SSE 关闭缓冲"],
            ["Docker Engine", "≥ 24.0", "容器运行时"],
            ["Docker Compose", "≥ v2.20", "容器编排"],
            ["Dify", "1.1.0（langgenius 官方镜像）", "Agent 编排 + RAG 引擎"],
            ["Python（构建基础镜像）", "3.11-slim-bookworm", "后端容器构建"],
            ["Node（构建基础镜像）", "20-alpine", "前端容器构建"],
            ["MinIO", "RELEASE.2024-08-29T01-40-52Z", "文件存储（二期）"],
        ],
        col_widths=[5.0, 3.0, 7.3],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_chapter_4(doc):
    add_h1(doc, "第四章  关键技术方案")
    add_body(doc,
        "本章对平台涉及的关键技术点做深入说明，涵盖鉴权、权限、Dify 集成、"
        "SSE 流式代理、数据模型、文件上传、加密导出等核心模块。"
        "实施阶段必须严格遵循本章设计，避免偏离导致的隐性缺陷。")

    add_h2(doc, "4.1 鉴权与权限模型")

    add_h3(doc, "4.1.1 身份层")
    add_body(doc, "身份层采用邮箱 + 密码（bcrypt 哈希）+ JWT 的经典组合，关键设计如下：")
    add_bullet(doc, "密码使用 passlib[bcrypt] 哈希存储，禁止明文；")
    add_bullet(doc, "JWT 包含两层：access token（15 分钟）+ refresh token（7 天）；")
    add_bullet(doc, "JWT 通过 httpOnly + Secure + SameSite=Strict cookie 传递，前端 JS 不可访问；")
    add_bullet(doc, "JWT 载荷：user_id、roles[]、dept_id、jti、exp；")
    add_bullet(doc, "Refresh token 入库（refresh_tokens 表），支持主动撤销与轮转；")
    add_bullet(doc, "CSRF 防护依赖 SameSite=Strict（基本盘）+ 敏感写接口 Origin / Referer 白名单校验。")

    add_quote(doc,
        "密钥拆分原则：JWT_SECRET 仅用于 JWT 签名，ENCRYPTION_KEY 仅用于 Dify API Key Fernet 加密。"
        "两个 key 完全独立，任一独立轮转不影响另一个。")

    add_h3(doc, "4.1.2 组织架构")
    add_body(doc, "组织架构采用三层模型：用户、部门、角色，多对多关系通过 user_roles 关联表实现。")
    add_data_table(doc,
        headers=["实体", "关键字段", "说明"],
        rows=[
            ["users", "id / email(unique) / name / password_hash / status / dept_id / timestamps",
             "status: 1=启用，0=禁用；dept_id 可空"],
            ["departments", "id / name / parent_id / path / created_at",
             "path: 物化路径 /1/3/7/，支持子树快速查询"],
            ["roles", "id / code(unique) / name",
             "内置 USER / APP_ADMIN / PLATFORM_ADMIN"],
            ["user_roles", "user_id + role_id（复合 PK）",
             "用户与角色多对多"],
        ],
        col_widths=[3.0, 6.5, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "4.1.3 应用授权（三态主体）")
    add_body(doc,
        "Dify 自带的\"应用\"在本平台统一称\"Agent\"。"
        "授权关系支持 user / dept / role 三种主体类型，用户对 Agent 的访问权限为：")

    add_quote(doc,
        "用户对 Agent 的访问权限 = 自己的授权 ∪ 所属部门的授权 ∪ 拥有角色的授权")

    add_body(doc, "授权解析函数伪代码：")
    add_code_block(doc, """
async def user_can_access_app(session, user_id, app_id) -> bool:
    # 1. 直接 user 授权
    if exists(user grant for app_id): return True
    # 2. 部门授权（user.dept_id）
    if user.dept_id and exists(dept grant for app_id, user.dept_id): return True
    # 3. 角色授权（user.roles[]）
    if exists(role grant for app_id, any_of(user.roles)): return True
    return False
    """, language="伪代码")

    add_body(doc, "FastAPI 侧必须在调用 Dify 前**再校验一次**权限，防止前端绕过。")

    add_h3(doc, "4.1.4 JWT 载荷与权限决策")
    add_body(doc,
        "JWT 载荷中已包含 user_id / roles[] / dept_id / jti，"
        "权限判断无需每次查询 DB，仅在敏感操作（如管理员后台）按需 reload 用户最新状态。"
        "jti 字段预留用于二期 token 黑名单 / 主动下线。")

    add_h3(doc, "4.1.5 管理员后台")
    add_body(doc, "PLATFORM_ADMIN 角色可执行以下操作：")
    add_bullet(doc, "用户 / 部门 / 角色 CRUD；")
    add_bullet(doc, "按 user / dept / role 维度给 Agent 授权；")
    add_bullet(doc, "查看审计日志 + 对话统计；")
    add_bullet(doc, "审计日志导出（CSV / JSON，二期补按条件过滤）。")

    add_h2(doc, "4.2 Dify 集成方案")

    add_h3(doc, "4.2.1 API Key 管理")
    add_body(doc,
        "每个 Dify 应用对应一个 API Key，存于 dify_api_keys 表的 api_key_encrypted 字段（Fernet 加密）。"
        "启动时解密加载到内存，运维通过管理接口改 Key 无需重启。"
        "加密密钥 ENCRYPTION_KEY 与 JWT 密钥 JWT_SECRET 独立存放。")

    add_h3(doc, "4.2.2 应用列表同步")
    add_body(doc, "应用同步策略如下：")
    add_bullet(doc, "启动时全量拉取 Dify 应用列表（GET /v1/apps），upsert 到本地 apps 表；")
    add_bullet(doc, "增量同步：MVP 阶段先用定时任务（每 5 分钟）跑全量，"
                 "二期升级为 Webhook（/internal/webhook/dify/app-events）；")
    add_bullet(doc, "同步时一并更新 dify_api_keys（应用 API Key）；")
    add_bullet(doc, "新增 / 删除应用通过 Webhook 通知 Web 后端。")

    add_h3(doc, "4.2.3 调用 Dify 应用")
    add_body(doc, "FastAPI 通过 httpx.AsyncClient 异步调用 Dify，关键设计：")
    add_bullet(doc, "客户端单例 + 连接池复用（lifespan 启动时构造，关闭时释放）；")
    add_bullet(doc, "超时：httpx.Timeout(120.0, connect=10.0)；")
    add_bullet(doc, "重试策略按调用类型区分：")
    add_quote(doc,
        "流式调用（chat-messages）：不可重试——重试会丢失已发出的 token，"
        "失败直接返回 event:error + 写审计；\n\n"
        "非流式调用（apps 列表 / files/upload / webhook 验证）："
        "指数退避（1s / 2s / 4s），最多 3 次；只重试 5xx 和网络错误，4xx 不重试。")

    add_h3(doc, "4.2.4 文件上传")
    add_body(doc, "文件上传链路：")
    add_code_block(doc, """
前端 → POST /api/chat/files (multipart) → FastAPI
    → 存本地 / MinIO
    → POST Dify /files/upload → 拿 dify_file_id
    → 存 messages.files 关联 dify_file_id
    """, language="流程示意")

    add_h3(doc, "4.2.5 Dify Webhook")
    add_body(doc, "Dify 通过 Webhook 回调 /internal/webhook/* 端点，处理两类事件：")
    add_bullet(doc, "应用 CRUD → 更新本地 apps 表；")
    add_bullet(doc, "工作流 / Agent 完成事件 → 兜底补齐 messages（防 SSE 中断漏写）。")

    add_h3(doc, "4.2.6 部署与隔离")
    add_body(doc, "Dify 部署关键点：")
    add_bullet(doc, "官方 docker-compose.yaml 独立部署；")
    add_bullet(doc, "生产环境外部化 Postgres / Redis；")
    add_bullet(doc, "子域名 dify.internal.company.com（内网访问）；")
    add_bullet(doc, "FastAPI 通过 DIFY_BASE_URL 环境变量连接。")

    add_h3(doc, "4.2.7 社区版限制与对策")
    add_data_table(doc,
        headers=["限制", "影响范围", "对策"],
        rows=[
            ["无原生多租户", "不影响——我们只装一个工作空间，\"多租户\"由 FastAPI 实现", "—"],
            ["应用级权限缺失", "仅影响 Dify 内 5–20 个管理员 / 编辑者",
             "编辑者按角色分组成员；普通员工由 FastAPI app_authorizations 授权"],
            ["不官方支持 SSO/MFA", "不影响——普通员工走 FastAPI JWT，不登录 Dify",
             "Dify 账号持有者少，靠管理流程约束"],
            ["共享工作空间", "离职编辑者可改坏所有 App",
             "运营流程约束 + 离职联动禁用（接口预留）"],
        ],
        col_widths=[3.5, 5.5, 6.3],
        first_col_bold=True,
    )

    add_quote(doc,
        "运营流程约束：Dify 账号最小化原则——仅给 5–20 人发 Dify 账号（管理员 + Agent 编辑者）；"
        "编辑者按小组隔离；任何 Dify App 配置变更都通过审计日志间接记录。")

    add_h2(doc, "4.3 对话流与 SSE 代理")

    add_h3(doc, "4.3.1 消息时序")
    add_code_block(doc, """
前端                FastAPI                  Dify
 │ POST /api/chat    │                       │
 │ {app_id, query,   │                       │
 │  conversation_id?,│                       │
 │  files?}          │                       │
 │─────────────────> │                       │
 │                   │ 1. JWT 鉴权           │
 │                   │ 2. 权限校验           │
 │                   │ 3. 限流（可选）       │
 │                   │ 4. 落库 user_msg      │
 │                   │ 5. 转发到 Dify        │
 │                   │    POST /chat-messages│
 │                   │─────────────────────> │
 │                   │<─────────────────────│
 │                   │   SSE chunks          │
 │   SSE: event:     │                       │
 │   message,        │                       │
 │   data:{...}      │                       │
 │<──────────────────│                       │
 │                   │                       │
 │   SSE: event:     │                       │
 │   message_end,    │                       │
 │   data:{...}      │                       │
 │<──────────────────│                       │
 │                   │ 6. 落库 assistant_msg │
 │                   │ 7. 审计日志           │
    """, language="时序图")

    add_h3(doc, "4.3.2 SSE 透传关键实现")
    add_body(doc,
        "SSE 透传必须避免任何形式的 buffer，否则用户需等待完整回复才能看到首字节。"
        "以下是反例与正例对比：")

    add_h4(doc, "反例（buffer）")
    add_code_block(doc, """
resp = await client.post(...)
full_body = await resp.aread()      # ❌ 阻塞等 Dify 全部生成完
return Response(full_body)          # 用户等 10 秒才看到整段回复
    """, language="Python")

    add_h4(doc, "正例（逐 chunk 透传 + 错误兜底）")
    add_code_block(doc, """
from fastapi.responses import StreamingResponse
import httpx

# 120s 总超时 / 10s 连接超时
DIFY_TIMEOUT = httpx.Timeout(120.0, connect=10.0)

async def _dify_event_generator(request_body, conv_id, app_id, user_id):
    accumulated_content = ""
    try:
        async with httpx.AsyncClient(timeout=DIFY_TIMEOUT) as client:
            async with client.stream(
                "POST",
                f"{DIFY_BASE}/v1/chat-messages",
                headers={"Authorization": f"Bearer {DIFY_KEY}"},
                json=request_body,
            ) as resp:
                async for line in resp.aiter_lines():
                    if not line:
                        continue
                    if line.startswith("data:") and '"event":"message"' in line:
                        try:
                            data = json.loads(line[5:].strip())
                            accumulated_content += data.get("answer", "")
                        except Exception:
                            pass
                    yield line + "\\n"     # ✅ 立即写给前端
    except httpx.TimeoutException:
        yield f"event: error\\ndata: {{\\"message\\": \\"Dify timeout\\"}}\\n\\n"
        await audit_log(user_id, "chat.timeout", conv_id, app_id)
    except Exception as e:
        yield f"event: error\\ndata: {{\\"message\\": \\"Proxy error\\"}}\\n\\n"
        await audit_log(user_id, "chat.error", conv_id, app_id, error=str(e))
    finally:
        if accumulated_content:
            await save_assistant_message(
                conversation_id=conv_id,
                content=accumulated_content,
                user_id=user_id,
                app_id=app_id,
            )
        yield "event: agent\\ndone\\ndata: {}\\n\\n"

@app.post("/api/chat/send")
async def send_message(body: ChatRequest, user = Depends(current_user)):
    return StreamingResponse(
        _dify_event_generator(dify_payload, conv.id, app.id, user.id),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",      # 关掉 nginx 缓冲
        },
    )
    """, language="Python")

    add_h3(doc, "4.3.3 易踩坑点")
    add_data_table(doc,
        headers=["坑", "现象", "解法"],
        rows=[
            ["await resp.aread()", "用户等到结束才看到字", "改用 aiter_lines() / aiter_bytes()"],
            ["uvicorn 默认缓冲", "大响应偶发卡顿", "header X-Accel-Buffering: no"],
            ["nginx proxy_buffering on", "同上", "nginx 配置 proxy_buffering off;"],
            ["前端用 EventSource", "断线自动重发整条消息", "改用 fetch + ReadableStream"],
            ["客户端 abort 但没 finally", "assistant 消息漏写", "必须 finally 块强制写库"],
        ],
        col_widths=[5.0, 4.5, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "4.3.4 Dify 流式事件类型")
    add_bullet(doc, "message：增量 token，逐字累加；")
    add_bullet(doc, "message_end：结束，含 metadata（usage / retriever_resources 等）；")
    add_bullet(doc, "error：异常事件；")
    add_bullet(doc, "message_file：图片等附件引用。")

    add_body(doc,
        "我们在 message_end 后**追加一个自有事件 agent_done**，"
        "告诉前端\"这轮结束，可落本地缓存 / 打点了\"。")

    add_h3(doc, "4.3.5 流中断与重连")
    add_bullet(doc, "鉴权：cookie 自动带 JWT（EventSource 不支持自定义 header）；")
    add_bullet(doc, "中断：MVP 不做断点续传，二期再实现 fetch 手动控制中断 + 重发机制；")
    add_bullet(doc, "兜底：客户端断流时 finally 块强制写 assistant 消息；")
    add_bullet(doc, "对账：Dify Webhook 兜底补齐 messages（防 SSE 中断漏写）。")

    add_h3(doc, "4.3.6 对话 ID")
    add_body(doc, "对话 ID 设计采用内外双 ID：")
    add_bullet(doc, "本平台生成的 conversations.id 使用 UUID，作为内部 API 主键；")
    add_bullet(doc, "Dify 的 conversation_id 存于 conversations.dify_conversation_id，做关联；")
    add_bullet(doc, "多轮对话：前端带内部 conv_id 来，FastAPI 取对应的 Dify conv_id 透传。")

    add_h2(doc, "4.4 数据模型设计")
    add_body(doc, "本节梳理核心数据表的关键字段与索引设计。")

    add_h3(doc, "4.4.1 数据表清单")
    add_code_block(doc, """
-- 用户 / 组织 / 角色
users (id, email UNIQUE, name, password_hash, status, dept_id, timestamps)
departments (id, name, parent_id, path, created_at)
roles (id, code UNIQUE, name)
user_roles (user_id, role_id) -- 复合 PK

-- Agent 镜像 + 授权
apps (id, dify_app_id UNIQUE, name, description, mode, status, synced_at, created_at)
app_authorizations (app_id, principal_type, principal_id)
  -- principal_type: 'user' | 'dept' | 'role'；复合 PK
  -- CheckConstraint: principal_type IN ('user','dept','role')

-- Dify 凭据
dify_api_keys (id, app_id UNIQUE, api_key_encrypted, updated_at)

-- 对话镜像
conversations (id UUID PK, user_id, app_id, dify_conversation_id,
               title, message_count, token_usage JSONB,
               created_at, updated_at, deleted_at)  -- 软删

messages (id, conversation_id UUID FK, role, content TEXT,
          dify_message_id, dedupe_key UNIQUE, files JSONB, created_at)

-- 审计
audit_logs (id, user_id, action, resource_type, resource_id,
            ip, user_agent, metadata JSONB, created_at)

-- Refresh Token
refresh_tokens (id, user_id, token_hash UNIQUE, expires_at,
                revoked_at, created_at)
    """, language="SQL")

    add_h3(doc, "4.4.2 关键决策与索引")
    add_data_table(doc,
        headers=["决策项", "理由"],
        rows=[
            ["conversations.id 用 UUID",
             "避免泄露平台用户规模，对外引用安全"],
            ["软删 deleted_at",
             "合规场景\"员工离职删账号\"必须保留审计痕迹"],
            ["messages.content 直接 TEXT",
             "MVP 不做向量化检索；跨对话检索二期再加 pgvector"],
            ["audit_logs 不存 message 原文",
             "按 resource_id join；日志表只存元数据，避免敏感信息二次泄露"],
            ["不引 Redis",
             "JWT 无状态；限流二期"],
            ["物化路径索引（path text_pattern_ops）",
             "加速子树查询（where path like '/1/3/%'）"],
            ["refresh_tokens 独立表",
             "支持主动撤销 + 轮转"],
            ["messages.dedupe_key UNIQUE 索引",
             "Webhook 兜底补写时防 SSE + Webhook 重复落库"],
        ],
        col_widths=[5.5, 9.8],
        first_col_bold=True,
    )

    add_h3(doc, "4.4.3 术语约定")
    add_quote(doc,
        "对外（前端 UI、文档）：统一称\"Agent\"；\n"
        "对内（DB、代码、API）：统一用 app / apps / app_id；\n"
        "这两个是同一个东西：apps 表里的一行 ≈ 前端展示的一个 Agent。")

    add_h2(doc, "4.5 文件上传方案")

    add_h3(doc, "4.5.1 安全约束")
    add_body(doc, "文件上传链路必须满足以下安全约束，缺一不可：")
    add_data_table(doc,
        headers=["约束", "规则", "实现位置"],
        rows=[
            ["大小上限", "≤ 20MB", "Content-Length 拒绝 + 流式累加双重"],
            ["MIME 白名单",
             "pdf / docx / txt / md / png / jpeg",
             "FastAPI 入参校验"],
            ["文件名清洗", "防路径穿越 / CRLF / 脚本注入",
             "werkzeug.utils.secure_filename"],
            ["审计", "每条上传记 audit_logs",
             "user_id / file_name / size / mime / sha256"],
            ["病毒扫描", "ClamAV", "二期再加"],
        ],
        col_widths=[2.8, 7.2, 5.3],
        first_col_bold=True,
    )

    add_h3(doc, "4.5.2 双层校验实现")
    add_code_block(doc, """
# Layer 1: Content-Length header 立即拒绝
content_length = request.headers.get("content-length")
if content_length and int(content_length) > 20 * 1024 * 1024:
    raise HTTPException(413, "File too large")

# Layer 2: 流式边读边累加
async for chunk in request.stream():
    bytes_read += len(chunk)
    if bytes_read > 20 * 1024 * 1024:
        # 主动断开连接
        await request.close()
        raise HTTPException(413, "File too large")
    """, language="Python")

    add_h3(doc, "4.5.3 存储位置")
    add_bullet(doc, "MVP 阶段：本地磁盘 app/storage/{yyyy}/{mm}/{file_id}.{ext}，文件名 UUID 防冲突；")
    add_bullet(doc, "二期：替换为 MinIO 对象存储。")

    add_h2(doc, "4.6 加密导出方案")

    add_h3(doc, "4.6.1 设计原则")
    add_data_table(doc,
        headers=["原则", "体现"],
        rows=[
            ["密码永不离开客户端", "派生、加密、解密全在浏览器 Web Crypto API；后端不存密码 / 派生密钥"],
            ["服务端不重复造轮子", "后端只负责流式返回原始数据，加密动作一律前端做"],
            ["离线可解", "提供 /decrypt 公共页面，员工不依赖任何 CLI 工具"],
            ["单一加密抽象", "所有下载共用一个 useEncryptedDownload hook"],
            ["可观测", "所有下载（含加密 / 非加密）都写 audit_logs"],
            ["可恢复的兼容扩展", "envelope 头带 v:1 版本字段；未来算法升级走 v2，不破坏旧文件"],
        ],
        col_widths=[3.5, 11.8],
        first_col_bold=True,
    )

    add_h3(doc, "4.6.2 加密文件格式（.enc）")
    add_code_block(doc, """
┌──────────────────────────────────────────────────────────────┐
│ 文件头（JSON，UTF-8，单行无换行）                              │
│ {                                                             │
│   "v": 1,                  // 格式版本号（当前仅支持 1）      │
│   "alg": "AES-256-GCM",    // 算法标识（allow-list）          │
│   "kdf": "PBKDF2-SHA256",  // 密钥派生函数（allow-list）      │
│   "iter": 100000,          // PBKDF2 迭代次数（100k–200k）    │
│   "salt": "<base64>",      // 16 字节随机盐                   │
│   "iv": "<base64>",        // 12 字节 GCM IV                  │
│   "ivLen": 12,             // 显式 IV 长度，防 Penumbra 升级  │
│   "authTag": "<base64>",   // 16 字节 GCM 认证标签            │
│   "origName": "conv-xxx.json",                                │
│   "mime": "application/json",                                 │
│   "createdAt": "2026-08-28T10:30:00Z"                         │
│ }                                                             │
├──────────────────────────────────────────────────────────────┤
│ \\n（一个换行符作为分隔符）                                    │
├──────────────────────────────────────────────────────────────┤
│ 加密数据（Penumbra 内部流，二进制）                            │
└──────────────────────────────────────────────────────────────┘
    """, language="Envelope")

    add_h3(doc, "4.6.3 envelope 验证 allow-list（/decrypt 强制）")
    add_data_table(doc,
        headers=["字段", "允许值", "拒绝行为"],
        rows=[
            ["v", "1（整数）", "v<1 / v>1 / 字符串 / NaN / 缺失 → 报错"],
            ["alg", "AES-256-GCM", "其他值 → 报错"],
            ["kdf", "PBKDF2-SHA256", "其他值 → 报错"],
            ["iter", "100000–200000", "范围外 → 报错（防 DoS）"],
            ["mime", "沿用上传白名单", "其他 → 警告但仍可解"],
            ["ivLen", "12（整数）", "其他 → 报错"],
            ["salt / iv / authTag", "长度正确 + base64 解码成功", "解码失败 → 报错"],
        ],
        col_widths=[3.5, 5.0, 6.8],
        first_col_bold=True,
    )

    add_h3(doc, "4.6.4 加密原语")
    add_data_table(doc,
        headers=["环节", "算法", "理由"],
        rows=[
            ["对称加密", "AES-256-GCM", "Web Crypto 原生；Penumbra 默认；机密性 + 完整性"],
            ["密钥派生", "PBKDF2-SHA256, 100,000 轮", "Web Crypto 原生；无需 WASM；UX 友好"],
            ["随机数", "crypto.getRandomValues()", "浏览器原生 CSPRNG"],
            ["密码强度评估", "zxcvbn", "DropBox 开源，纯前端，无网络调用"],
        ],
        col_widths=[3.0, 5.5, 6.8],
        first_col_bold=True,
    )

    add_h3(doc, "4.6.5 密码强度要求")
    add_bullet(doc, "长度 ≥ 12 字符（不再用 8）；")
    add_bullet(doc, "至少包含 2 类字符（小写 / 大写 / 数字 / 符号）；")
    add_bullet(doc, "zxcvbn 分数 ≥ 2 允许通过；分数 < 2 时强制显示警告；")
    add_bullet(doc, "任何分数 ≤ 2 都弹红字警告\"该密码可在离线攻击下被破解\"；")
    add_bullet(doc, "绝不强制高强度——员工会忘，但弱密码必须给警告。")

    add_h3(doc, "4.6.6 明确不做密码找回 / 密钥托管")
    add_quote(doc,
        "忘了密码 = 文件解不开。这是产品决策，写进：\n"
        "- PasswordDialog 顶部提示\n"
        "- /decrypt 页面顶部红色警告\n"
        "- 用户首次使用功能时的 onboarding tooltip")

    add_h3(doc, "4.6.7 后端导出端点")
    add_data_table(doc,
        headers=["端点", "方法", "用途"],
        rows=[
            ["/api/conversations/{id}/export?format=json|md", "GET", "单会话导出"],
            ["/api/conversations/export-batch", "POST", "批量会话导出（zip，≤ 100 个）"],
            ["/api/admin/audit-logs/export", "POST", "审计日志导出（CSV / JSON）"],
            ["/api/files/{file_id}/download", "GET", "附件下载"],
        ],
        col_widths=[7.0, 2.0, 6.3],
        first_col_bold=True,
    )

    add_h3(doc, "4.6.8 流式审计三态保证")
    add_body(doc, "所有流式下载端点统一采用三态审计模式，保证不漏写：")
    add_bullet(doc, "客户端正常完成 → BackgroundTask.on_success → audit completed=true；")
    add_bullet(doc, "客户端 abort → CancelledError 在 generator 捕获 → audit completed=false, reason=client_disconnect；")
    add_bullet(doc, "服务端异常 → Exception 捕获 → audit completed=false, reason=<ExceptionType>；")
    add_bullet(doc, "audit_written flag 保证三种路径互斥，恰好写一次。")

    add_h3(doc, "4.6.9 CSV 注入防护")
    add_code_block(doc, """
def sanitize_csv_cell(value: str) -> str:
    if not isinstance(value, str):
        value = str(value)
    if value and value[0] in ('=', '+', '-', '@', '\\t', '\\r'):
        value = "'" + value     # 前缀单引号防公式注入
    return value
    """, language="Python")

    add_h3(doc, "4.6.10 文件大小上限")
    add_data_table(doc,
        headers=["层", "限制"],
        rows=[
            ["nginx", "client_max_body_size 5g"],
            ["FastAPI", "StreamingResponse 内部累加 bytes_sent，超过 5 GB 主动断开"],
            ["前端", "<DownloadButton> 下载前 GET HEAD 检查 Content-Length，> 500 MB 拒绝"],
        ],
        col_widths=[3.0, 12.3],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_chapter_5(doc):
    add_h1(doc, "第五章  部署架构")

    add_h2(doc, "5.1 网络拓扑")
    add_code_block(doc, """
企业内网
   │
   ▼
浏览器 (员工电脑)
   │
   ▼
Nginx 反向代理 / HTTPS
   │ portal.internal.company.com
   ├─ /api/*  →  FastAPI
   └─ /*      →  前端静态
   │
   ▼
FastAPI (2 实例)
   │
   ├──> Postgres 16
   │
   └──> Dify (独立部署)
            │
            └──> 独立 Postgres / Redis
    """, language="网络拓扑")

    add_h2(doc, "5.2 网络隔离分档")
    add_data_table(doc,
        headers=["档位", "措施", "适用"],
        rows=[
            ["基础档", "Dify 公网 / 内网可达，靠 Dify 账号体系约束", "MVP（推荐）"],
            ["加固档", "加网络 ACL：仅 FastAPI 所在网段访问 Dify", "中型合规企业"],
            ["极致档", "Dify 离线，仅堡垒机访问", "金融 / 政府"],
        ],
        col_widths=[2.5, 8.0, 4.8],
        first_col_bold=True,
    )

    add_h2(doc, "5.3 Docker 编排")
    add_body(doc, "本项目 deploy/ 目录结构：")
    add_code_block(doc, """
deploy/
├── docker-compose.yml          # 主编排（Postgres + Nginx + 后端 + 前端）
├── docker-compose.override.yml # 本地开发覆盖（dev 配置、挂载代码）
├── .env.example
├── nginx/
│   └── conf.d/portal.conf
├── dify/                       # Dify 官方 compose（独立部署）
│   ├── docker-compose.yaml
│   └── .env.example
└── scripts/
    ├── init-dify.sh            # 初始化 Dify 管理员
    └── seed-admin.sh           # 创建 PLATFORM_ADMIN
    """, language="目录结构")

    add_body(doc,
        "为什么不把 Dify 放进本项目 compose？\n"
        "- Dify 官方 compose 自带 Postgres + Redis + 5 个微服务（api / worker / web / nginx / SSRF proxy），耦合度高；\n"
        "- 升级 Dify 时只想替换它的目录，不希望牵连本项目；\n"
        "- 本项目开发者经常需要单独重启 FastAPI，不希望连带重启 Dify。")

    add_h2(doc, "5.4 启动顺序")
    add_number(doc, "Postgres（独立实例）：docker compose up -d postgres，等待 healthcheck 通过；")
    add_number(doc, "Dify（官方 compose，外部 Postgres）：docker compose -f docker-compose.yaml up -d，等待 1–2 分钟；")
    add_number(doc, "初始化 Dify 管理员：浏览器访问 /install 完成初始化，或调用 /v1/setup/init-admin API；")
    add_number(doc, "FastAPI 跑迁移：alembic upgrade head，触发 Dify 应用同步初始化 apps 表；")
    add_number(doc, "Nginx 启动：门户可用。")

    add_h2(doc, "5.5 关键配置")

    add_h3(doc, "5.5.1 nginx.conf")
    add_code_block(doc, """
upstream backend {
    server backend:8000;
}

server {
    listen 80;
    server_name _;

    # SSE 必需：禁用缓冲
    proxy_buffering off;
    proxy_cache off;
    proxy_read_timeout 300s;

    location /api/ {
        proxy_pass http://backend;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;

        # SSE 特殊 header
        proxy_set_header Connection '';
        proxy_http_version 1.1;
        chunked_transfer_encoding off;
    }

    location /decrypt {
        add_header Content-Security-Policy "default-src 'self'; script-src 'self'; worker-src 'self' blob:; style-src 'self' 'unsafe-inline'" always;
    }

    location / {
        root /usr/share/nginx/html;
        try_files $uri $uri/ /index.html;
    }
}
    """, language="nginx.conf")

    add_h3(doc, "5.5.2 backend/.env 必备")
    add_code_block(doc, """
# 必须：JWT 签名密钥（≥ 32 字节随机）
JWT_SECRET=<python -c "import secrets; print(secrets.token_urlsafe(32))">

# 必须：Dify API Key 加密密钥（Fernet 格式）
ENCRYPTION_KEY=<python -c "from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())">

# Postgres
DATABASE_URL=postgresql+asyncpg://agent:agent@localhost:5432/agent_platform

# Dify
DIFY_BASE_URL=http://localhost:5001
DIFY_ADMIN_EMAIL=admin@dify.local
DIFY_ADMIN_PASSWORD=change-me-dify-admin-password

# 种子管理员
ADMIN_EMAIL=admin@company.local
ADMIN_PASSWORD=AdminPass123
ADMIN_NAME=平台管理员

# 开发模式
DEBUG=true
LOG_LEVEL=INFO
ALLOWED_ORIGINS=http://localhost:5173,http://127.0.0.1:5173
    """, language="Environment")

    add_quote(doc,
        "安全警告：上述密码仅用于本地开发。生产环境必须：\n"
        "- JWT_SECRET ≥ 64 字节随机；\n"
        "- ENCRYPTION_KEY 用 KMS 管理；\n"
        "- 所有密码不进 git，用 secrets 管理平台。")

    add_h3(doc, "5.5.3 CORS 策略")
    add_bullet(doc, "生产：Nginx 同源反代 → 后端无需 CORS 头；")
    add_bullet(doc, "开发：Vite 5173 → FastAPI 8000 跨域 → FastAPI 启用 CORS 中间件，"
                 "allow_origins 读取 ALLOWED_ORIGINS；")
    add_bullet(doc, "生产环境 ALLOWED_ORIGINS 仅含生产域名（如 https://portal.internal.company.com）；"
                 "DEBUG=true 时追加 localhost:5173 / 127.0.0.1:5173。")

    add_h3(doc, "5.5.4 FastAPI 多实例与 SSE")
    add_body(doc, "FastAPI 多实例部署关键点：")
    add_bullet(doc, "\"实例数 × workers\" = 容器副本数 × uvicorn worker 数（推荐 2 × 2 = 4 workers）；")
    add_bullet(doc, "SSE 连接由 nginx 随机分到任一 worker，**这是 OK 的**（SSE 长连接本身无状态）；")
    add_bullet(doc, "每个 worker 各自维护自己的 httpx AsyncClient 单例（**不要跨 worker 共享**）；")
    add_bullet(doc, "排障：每个 worker 打 worker_id 到日志，便于定位是哪个 worker 处理了某次 SSE。")

    add_h3(doc, "5.5.5 端口规划")
    add_data_table(doc,
        headers=["端口", "用途", "容器 / 进程"],
        rows=[
            ["5432", "Postgres", "docker: ap-postgres"],
            ["5001", "Dify API", "docker: dify-api"],
            ["3000", "Dify Web UI", "docker: dify-web"],
            ["6379", "Dify 内置 Redis", "docker: dify-redis"],
            ["8000", "FastAPI 后端（dev）", "本机 uvicorn"],
            ["5173", "Vite 前端（dev）", "本机 vite dev"],
            ["80 / 443", "Nginx（生产）", "docker: ap-nginx"],
        ],
        col_widths=[2.5, 6.0, 6.8],
        first_col_bold=True,
    )

    add_h3(doc, "5.5.6 镜像版本核对清单")
    add_body(doc, "启动前必查的 6 个版本，避免日后\"为什么我这跑不通\"的争议：")
    add_number(doc, "Postgres 16.x → docker run --rm postgres:16-alpine postgres --version")
    add_number(doc, "Redis 7.x（仅 Dify 用） → docker run --rm redis:7-alpine redis-server --version")
    add_number(doc, "Python 3.11.x → docker run --rm python:3.11-slim-bookworm python --version")
    add_number(doc, "Node 20.x → docker run --rm node:20-alpine node --version")
    add_number(doc, "Nginx 1.25.x → docker run --rm nginx:1.25-alpine nginx -v")
    add_number(doc, "Dify API 1.1.0 → docker run --rm --entrypoint echo langgenius/dify-api:1.1.0 \"Dify API 1.1.0\"")

    add_page_break(doc)


def write_chapter_6(doc):
    add_h1(doc, "第六章  安全方案")

    add_h2(doc, "6.1 身份认证安全")
    add_body(doc, "身份认证采用经典三件套：邮箱 + 密码 + JWT，叠加多项安全措施：")
    add_data_table(doc,
        headers=["措施", "说明"],
        rows=[
            ["bcrypt 哈希", "passlib[bcrypt]，禁止明文存储密码"],
            ["JWT 双 token", "access 15min + refresh 7d；refresh 入库支持主动撤销"],
            ["httpOnly + Secure + SameSite=Strict cookie", "前端 JS 不可访问；生产强制 HTTPS；防跨站 CSRF"],
            ["JWT 载荷 jti", "预留二期 token 黑名单 / 主动下线"],
            ["Refresh token 轮转", "旧 refresh 二次使用立即撤销并返回 401"],
            ["密钥拆分", "JWT_SECRET 与 ENCRYPTION_KEY 独立存放，独立轮转"],
            ["OAuth 接口预留", "/api/auth/oauth/{provider}/callback，二期接企业 SSO"],
        ],
        col_widths=[5.0, 10.3],
        first_col_bold=True,
    )

    add_h2(doc, "6.2 CSRF 防护")

    add_h3(doc, "6.2.1 双层防护")
    add_bullet(doc, "基本盘：SameSite=Strict（cookie 不随跨站请求发送）；")
    add_bullet(doc, "加固层：CSRFMiddleware 校验敏感写接口的 Origin / Referer 白名单。")

    add_h3(doc, "6.2.2 最小端点清单（Origin 必须校验）")
    add_data_table(doc,
        headers=["端点", "方法", "必须校验 Origin"],
        rows=[
            ["/api/chat/send", "POST", "✅"],
            ["/api/chat/files", "POST", "✅"],
            ["/api/auth/refresh", "POST", "✅"],
            ["/api/auth/logout", "POST", "✅"],
            ["/api/admin/*", "任意写", "✅"],
            ["GET / HEAD / OPTIONS", "任意", "❌ 跳过"],
        ],
        col_widths=[5.0, 3.0, 7.3],
        first_col_bold=True,
    )

    add_h3(doc, "6.2.3 中间件实现")
    add_code_block(doc, """
WRITE_METHODS = {"POST", "PUT", "PATCH", "DELETE"}
PROTECTED_PATHS = ("/api/chat/send", "/api/chat/files", "/api/auth/", "/api/admin/")

class CSRFMiddleware(BaseHTTPMiddleware):
    def __init__(self, app, allowed_origins: list[str]):
        super().__init__(app)
        self.allowed = {o.rstrip("/") for o in allowed_origins}

    async def dispatch(self, request, call_next):
        if request.method in WRITE_METHODS and \\
           any(request.url.path.startswith(p) for p in PROTECTED_PATHS):
            origin = request.headers.get("origin") \\
                     or _extract_origin_from_referer(request.headers.get("referer"))
            origin = (origin or "").rstrip("/")
            if not origin or origin not in self.allowed:
                return JSONResponse({"detail": "Forbidden: invalid origin"}, status_code=403)
        return await call_next(request)
    """, language="Python")

    add_h2(doc, "6.3 文件上传安全")
    add_body(doc, "文件上传链路必须满足三重校验 + 一道审计：")
    add_data_table(doc,
        headers=["维度", "规则", "实现"],
        rows=[
            ["大小", "≤ 20MB",
             "Content-Length header 立即拒绝 + 流式累加双层校验"],
            ["MIME", "白名单（pdf / docx / txt / md / png / jpeg）",
             "FastAPI 入参校验"],
            ["文件名", "防路径穿越 / CRLF / 脚本注入",
             "werkzeug.utils.secure_filename"],
            ["审计", "记录每条上传的 user_id / file_name / size / mime / sha256",
             "audit_logs INSERT"],
        ],
        col_widths=[2.5, 8.0, 4.8],
        first_col_bold=True,
    )

    add_quote(doc,
        "二期再加 ClamAV 病毒扫描 + Sensitive Word 扫描。\n"
        "当前 MVP 阶段依赖白名单 + 大小限制做基本防御。")

    add_h2(doc, "6.4 数据加密")
    add_body(doc, "本平台数据加密分两层：")
    add_data_table(doc,
        headers=["层级", "场景", "算法", "密钥管理"],
        rows=[
            ["存储加密", "Dify API Key 加密存储",
             "Fernet（AES-128-CBC + HMAC-SHA256）", "ENCRYPTION_KEY 环境变量，独立于 JWT_SECRET"],
            ["传输加密", "前端 ↔ 后端 ↔ Dify 全链路",
             "TLS 1.2+", "Nginx 证书 + 公司 PKI"],
            ["用户数据加密", "导出文件加密",
             "AES-256-GCM（浏览器侧）", "用户密码派生，PBKDF2 100k 轮"],
        ],
        col_widths=[3.0, 4.0, 4.5, 3.8],
        first_col_bold=True,
    )

    add_h3(doc, "6.4.1 密钥管理红线")
    add_bullet(doc, "JWT_SECRET ≥ 32 字节随机（生产 ≥ 64 字节）；")
    add_bullet(doc, "ENCRYPTION_KEY 用 KMS 管理（开发环境可用 Fernet.generate_key() 生成）；")
    add_bullet(doc, "所有密码不进 git，用 secrets 管理平台；")
    add_bullet(doc, "用户密码永不出客户端（加密导出场景）；")
    add_bullet(doc, "派生 key 用完立刻 fill(0) 清内存。")

    add_h2(doc, "6.5 审计日志")

    add_h3(doc, "6.5.1 审计原则")
    add_bullet(doc, "覆盖所有敏感操作（登录、对话发起、文件上传、Agent 授权变更、对话导出等）；")
    add_bullet(doc, "audit_logs 不存 message 原文，只存元数据 + resource_id（按需 join）；")
    add_bullet(doc, "保留 ≥ 180 天；")
    add_bullet(doc, "支持按时间窗口 + 用户 + 动作类型多维查询；")
    add_bullet(doc, "支持管理员导出 CSV / JSON（带公式注入防护）。")

    add_h3(doc, "6.5.2 审计动作枚举")
    add_data_table(doc,
        headers=["动作", "触发场景", "元数据关键字段"],
        rows=[
            ["auth.login", "登录成功", "email, ip, user_agent"],
            ["auth.logout", "注销", "ip, user_agent"],
            ["auth.refresh", "refresh 轮转", "ip, user_agent"],
            ["chat.send", "发起对话", "conv_id, app_id, ip, duration_ms"],
            ["chat.timeout", "Dify 调用超时", "conv_id, app_id, timeout"],
            ["chat.error", "Dify 调用异常", "conv_id, app_id, error"],
            ["file.upload", "文件上传", "file_name, size, mime, sha256"],
            ["file.download", "附件下载", "file_id, name, size, mime, bytes_sent"],
            ["conversation.export", "单会话导出", "format, bytes, completed"],
            ["conversation.export_batch", "批量导出", "batch_id, format, requested, processed, bytes"],
            ["audit_logs.export", "审计导出", "from, to, format, rows, bytes"],
            ["admin.update_user", "用户信息变更", "user_id, changes"],
            ["admin.grant_authorization", "Agent 授权变更", "app_id, principal_type, principal_id"],
            ["user.terminate", "离职联动禁用", "user_id, action_taken"],
        ],
        col_widths=[4.0, 5.0, 6.3],
        first_col_bold=True,
    )

    add_h3(doc, "6.5.3 流式审计三态保证")
    add_body(doc, "所有流式端点（对话、文件下载、导出）必须实现三态审计：")
    add_number(doc, "正常完成 → BackgroundTask.on_success → audit completed=true")
    add_number(doc, "客户端 abort → CancelledError 在 generator 捕获 → audit completed=false, reason=client_disconnect")
    add_number(doc, "服务端异常 → Exception 捕获 → audit completed=false, reason=<ExceptionType>")
    add_number(doc, "audit_written flag 保证三种路径互斥，恰好写一次")

    add_h3(doc, "6.5.4 日志脱敏红线")
    add_bullet(doc, "loguru 严禁打印完整 message content / 用户密码 / Dify API Key；")
    add_bullet(doc, "audit_logs.metadata 不存 message 原文；")
    add_bullet(doc, "错误日志中 token / cookie 一律脱敏为 ***；")
    add_bullet(doc, "定期 grep \"password\" / \"secret\" 关键词做审计。")

    add_page_break(doc)


def write_chapter_7(doc):
    add_h1(doc, "第七章  测试方案")

    add_h2(doc, "7.1 测试策略")
    add_body(doc, "本项目采用三层测试金字塔 + E2E 兜底的策略：")
    add_quote(doc,
        "单元测试（Vitest / pytest）→ 集成测试（pytest）→ E2E（Playwright）→ 手动验证清单\n\n"
        "关键模块单测覆盖率 ≥ 70%；E2E 覆盖关键业务路径；手动验证补自动化盲区。")

    add_h2(doc, "7.2 单元测试")

    add_h3(doc, "7.2.1 后端（pytest + httpx AsyncClient）")
    add_body(doc, "关键模块单测清单：")
    add_data_table(doc,
        headers=["模块", "用例", "验证点"],
        rows=[
            ["app.dify", "DifyClient.list_apps / upload_file / chat_messages_stream",
             "Mock Dify API；覆盖正常路径 + 5xx 重试 + 流式 chunks"],
            ["app.chat", "_dify_event_generator",
             "fake async generator 验证透传 + finally 块写库 + 中断兜底"],
            ["app.apps.authorization", "user_can_access_app",
             "覆盖 user / dept / role 三态授权 + 复合授权 + 无授权"],
            ["app.auth.service", "authenticate / issue_tokens / rotate_refresh_token",
             "覆盖密码错误 / refresh 轮转 / 用户禁用"],
            ["app.core.security", "JWT 签发 + 校验 + refresh hash",
             "覆盖过期 token / 篡改 token / jti 唯一"],
            ["app.utils.filename", "sanitize_filename",
             "覆盖路径穿越 / CRLF / bidi override / 超长名"],
        ],
        col_widths=[4.0, 5.5, 5.8],
        first_col_bold=True,
    )

    add_h3(doc, "7.2.2 前端（Vitest + React Testing Library）")
    add_data_table(doc,
        headers=["模块", "用例", "验证点"],
        rows=[
            ["lib/crypto/envelope", "wrap → unwrap 往返",
             "envelope 解析正确 + 所有字段一致"],
            ["lib/crypto/envelope", "拒绝非法 alg / kdf / iter / v / ivLen",
             "allow-list 校验"],
            ["lib/crypto/envelope", "拒绝 base64 解码失败",
             "salt / iv / authTag 解码校验"],
            ["lib/crypto/deriveKey", "相同密码 + salt + iter → 相同 key",
             "派生函数决定性"],
            ["lib/crypto/deriveKey", "不同 salt → 不同 key",
             "派生函数隔离性"],
            ["lib/crypto/encryptStream", "加密 → 解密 → 原文一致",
             "端到端"],
            ["lib/crypto/decryptFile", "错误密码 → 抛错",
             "GCM authTag 校验失败"],
            ["lib/crypto/sanitizeFilename", "../etc/passwd / CRLF / bidi",
             "路径穿越 / CRLF 注入 / bidi override"],
        ],
        col_widths=[4.5, 5.5, 5.3],
        first_col_bold=True,
    )

    add_h2(doc, "7.3 集成测试（pytest）")
    add_body(doc, "覆盖所有流式端点 + 关键鉴权 / 授权路径：")
    add_data_table(doc,
        headers=["端点", "用例", "验证点"],
        rows=[
            ["/api/conversations/{id}/export?format=json", "流式返回",
             "流式分块正确 + 字节数对得上 + JSON 解析通过"],
            ["/api/conversations/{id}/export?format=md", "Markdown 格式",
             "Markdown 格式正确"],
            ["/api/conversations/{id}/export", "跨用户访问",
             "用户 A 不能导出用户 B 的对话（403）"],
            ["/api/conversations/{id}/export", "客户端中途断开",
             "audit 包含 completed=false（finally 块生效）"],
            ["/api/conversations/export-batch", "批量导出 100 个对话",
             "生成有效 zip + 行数 == 100"],
            ["/api/conversations/export-batch", "批量 > 100",
             "返回 400（上限校验）"],
            ["/api/files/{file_id}/download", "流式返回",
             "MD5 一致"],
            ["/api/admin/audit-logs/export?format=csv", "CSV 注入防护",
             "构造 action=cmd → 导出后前缀 ' 防公式注入"],
            ["/api/admin/audit-logs/export?format=csv", "CSV 转义",
             "含逗号 / 引号 / 换行的字段正确转义"],
        ],
        col_widths=[5.5, 4.0, 5.8],
        first_col_bold=True,
    )

    add_h2(doc, "7.4 E2E 测试（Playwright）")
    add_data_table(doc,
        headers=["路径", "验证点"],
        rows=[
            ["登录 → 选 Agent → 发起对话 → 看到流式回复 → 历史里能看到",
             "完整闭环"],
            ["权限拒绝路径：用户试图调未授权 Agent", "403 + 审计落库"],
            ["Dify 5xx 路径：mock Dify 返回 500",
             "前端看到 event:error + 审计落库"],
            ["文件上传失败路径：超大文件 / 非法 MIME",
             "前端 + 后端双重拒绝"],
            ["JWT 过期刷新路径：access 过期自动用 refresh 续期",
             "无感刷新"],
            ["加密下载 → /decrypt 解密 → 拿到原文",
             "双向一致"],
            ["/decrypt 上传篡改了 alg 字段的 .enc → 报错",
             "allow-list 校验"],
            ["弱密码（zxcvbn=1）→ 弹警告但可继续",
             "警告 UI"],
        ],
        col_widths=[8.0, 7.3],
        first_col_bold=True,
    )

    add_h2(doc, "7.5 手动验证清单")
    add_body(doc, "以下场景必须人工验证，自动化难以覆盖：")
    add_number(doc, "SSE 首字节延迟 ≤ 300ms；")
    add_number(doc, "断流后页面不崩（finally 块正常触发）；")
    add_number(doc, "删除用户联动禁用 Dify；")
    add_number(doc, "Chrome / Edge / Firefox 三浏览器加密导出全流程可用；")
    add_number(doc, "DevTools Memory snapshot 确认派生 key 不残留；")
    add_number(doc, "DevTools Network 确认 password 不上行。")

    add_page_break(doc)


def write_chapter_8(doc):
    add_h1(doc, "第八章  项目实施计划")

    add_h2(doc, "8.1 阶段划分")
    add_body(doc,
        "本项目 MVP 阶段预计 6–10 周，分 8 个阶段交付。"
        "原计划 4–6 周过于乐观——Dify 生产部署、外部化 PG / Redis、SSE 三层联调、"
        "多态授权一致性、CSRF / 文件上传安全均需工时。")

    add_data_table(doc,
        headers=["阶段", "内容", "交付物", "估时"],
        rows=[
            ["0", "仓库骨架与基础设施",
             "仓库结构 + Docker Compose 起 Dify + Postgres", "1 周"],
            ["1", "后端鉴权与权限模型",
             "用户登录 / JWT / CSRF 可用，E2E 通过", "1.5 周"],
            ["2", "Dify 适配层 + Agent 同步 + 授权",
             "可列出已授权 Agent，权限校验生效", "1 周"],
            ["3", "对话代理 + SSE + 对话镜像",
             "API 调通可流式对话，本地落库", "2 周"],
            ["4", "文件上传（带安全约束）",
             "文件上传全链路可用", "0.5 周"],
            ["5", "前端骨架 + 登录 + 对话 UI",
             "用户端 MVP 可用", "2 周"],
            ["6", "管理员后台",
             "管理员可管用户 / 部门 / Agent 授权", "1 周"],
            ["7", "生产部署 + 可观测性",
             "生产环境可用，监控告警就位", "1 周"],
            ["合计", "—", "—", "10 周"],
        ],
        col_widths=[1.5, 5.0, 5.5, 3.3],
        first_col_bold=True,
    )

    add_quote(doc,
        "本文档详细描述阶段 0、1、2、3、4。阶段 5–7 建议作为后续独立 plan 文件实施。")

    add_h2(doc, "8.2 资源投入")
    add_data_table(doc,
        headers=["角色", "人数", "职责"],
        rows=[
            ["后端工程师", "2", "FastAPI 后端开发 + Dify 集成 + 数据库设计"],
            ["前端工程师", "1", "React SPA + 对话 UI + 管理员后台 + 加密导出"],
            ["SRE / DevOps", "0.3（兼）", "Docker 编排 + Nginx 配置 + 监控告警"],
            ["安全 review", "0.2（兼）", "鉴权 / CSRF / 文件上传 / 加密导出的安全评审"],
            ["产品 / 测试", "0.5（兼）", "需求确认 + E2E 测试 + 用户验收"],
            ["合计", "约 4 人 × 10 周", "—"],
        ],
        col_widths=[3.0, 3.0, 9.3],
        first_col_bold=True,
    )

    add_h2(doc, "8.3 里程碑")
    add_data_table(doc,
        headers=["里程碑", "时间节点", "标志事件"],
        rows=[
            ["M1 - 环境就绪", "第 1 周末", "Postgres + Dify 起 + alembic 迁移成功 + 健康检查通过"],
            ["M2 - 鉴权可用", "第 2.5 周末", "登录 + JWT + CSRF 单测通过 + E2E 跑通"],
            ["M3 - 授权生效", "第 3.5 周末", "三态授权解析 + 权限校验单测覆盖"],
            ["M4 - 对话可用", "第 5.5 周末", "SSE 流式对话 + 落库 + finally 兜底"],
            ["M5 - 文件上传", "第 6 周末", "20MB + MIME 白名单 + 文件名清洗全链路"],
            ["M6 - 前端 MVP", "第 8 周末", "用户可登录 + 对话 + 历史 + 加密导出"],
            ["M7 - 管理后台", "第 9 周末", "PLATFORM_ADMIN 可管用户 / Agent 授权"],
            ["M8 - 上线", "第 10 周末", "生产部署 + 监控告警就位 + 验收通过"],
        ],
        col_widths=[4.0, 3.0, 8.3],
        first_col_bold=True,
    )

    add_page_break(doc)


def write_chapter_9(doc):
    add_h1(doc, "第九章  风险与对策")

    add_h2(doc, "9.1 技术风险")
    add_data_table(doc,
        headers=["风险", "严重度", "缓解"],
        rows=[
            ["SSE 缓冲导致体感差", "高",
             "header + 代码双重保险（X-Accel-Buffering:no + proxy_buffering off + aiter_lines）；上线前压测首字节延迟"],
            ["Dify API 变更破坏适配层", "中",
             "锁版本（langgenius/dify-api:1.1.0），升级有专项 review"],
            ["对话数据双写不一致", "中",
             "Webhook 兜底同步 + 定期对账脚本（仅告警不删）+ messages.dedupe_key UNIQUE 索引"],
            ["LLM 成本失控", "中",
             "二期引入 token 配额（每用户 / 部门每月上限）"],
            ["前端 SSE 断线导致内容丢失", "中",
             "MVP 不做断点续传；二期实现 fetch 手动中断 + 重发；当前 finally 块保证 assistant 消息必落库"],
        ],
        col_widths=[4.0, 1.5, 9.8],
        first_col_bold=True,
    )

    add_h2(doc, "9.2 进度风险")
    add_data_table(doc,
        headers=["风险", "严重度", "缓解"],
        rows=[
            ["Dify 部署调通延期", "高",
             "阶段 0 提前做 Dify 部署验证；外部化 Postgres / Redis 优先；镜像版本锁定"],
            ["SSE 三层联调调试周期长", "中",
             "nginx + uvicorn + httpx 三层同时调；预置调试工具（worker_id 日志）；压测脚本提前准备"],
            ["多态授权一致性单测覆盖不全", "中",
             "阶段 2 强制 user_can_access_app 100% 覆盖率；复合授权（user + dept + role 同时）单测"],
            ["加密导出浏览器兼容性问题", "中",
             "目标 Chrome / Edge / Firefox ≥ 102；明确不支持 Safari；上线前跨浏览器验证"],
        ],
        col_widths=[5.0, 1.5, 8.8],
        first_col_bold=True,
    )

    add_h2(doc, "9.3 安全风险")
    add_data_table(doc,
        headers=["风险", "严重度", "缓解"],
        rows=[
            ["离职员工用旧 cookie 调通", "中",
             "JWT 短过期（15min）+ jti 黑名单机制（二期）；refresh 主动撤销；离职触发软删对话"],
            ["CSRF / 文件上传安全漏洞", "高",
             "SameSite=Strict + Origin 校验 + 文件白名单 + 大小限制；e2e 测试覆盖异常路径"],
            ["SECRET_KEY 单点泄漏", "中",
             "已拆分为 JWT_SECRET + ENCRYPTION_KEY（见 4.1.1）；独立轮转；KMS 管理"],
            ["弱密码导致 .enc 可破解", "中",
             "zxcvbn 警告 + 12 字符最小 + 2 类字符；员工遗忘密码风险已在 /decrypt 提示"],
            ["CSV 公式注入", "中",
             "Python csv 模块 + sanitize_csv_cell 前缀 = + - @ 防公式注入"],
            ["origName 文件名注入", "低",
             "sanitize_filename 服务端 + 前端双重清洗；防路径穿越 / CRLF / bidi override"],
            ["envelope header 字段值未校验", "中",
             "/decrypt 强制 allow-list（alg / kdf / iter 范围 / ivLen / version）"],
            ["审计日志泄露用户下载内容", "低",
             "audit_logs 只存元数据，不存原文；导出流式元数据即可"],
            ["暴力破解 .enc", "中",
             "PBKDF2 100k 轮降低破解速度；强密码是唯一防线"],
        ],
        col_widths=[5.0, 1.5, 8.8],
        first_col_bold=True,
    )

    add_h2(doc, "9.4 待办事项")
    add_body(doc, "以下事项需在实施初期对齐，避免后期返工：")
    add_number(doc, "选定部署机器规格（FastAPI 2 实例至少 2C4G，"
                    "Dify 至少 8C16G——官方镜像内嵌 Postgres + Redis 很重，生产必须外部化）；")
    add_number(doc, "选定 Dify 镜像版本（锁定次版本号，定期升级）；")
    add_number(doc, "Dify API Key 轮转策略（建议 90 天）；")
    add_number(doc, "离职联动禁用 Dify 账号的实现路径确认："
                    "优先调研 Dify 社区版\"禁用账号\"管理 API 是否真的存在（查 Swagger）；"
                    "备选若 API 不存在或权限不够，降级为\"手动禁用脚本 + 审计触发提醒\"；")
    add_number(doc, "日志脱敏清单：loguru 严禁打印完整 message content / 用户密码 / Dify API Key；"
                    "audit_logs.metadata 不存 message 原文；错误日志中 token / cookie 一律脱敏为 ***；")
    add_number(doc, "对账脚本设计：每天凌晨基于 dify_app_id 集合 diff，"
                    "丢失的本地授权关系通过管理员告警而不是自动删除（避免脚本误删授权）；")
    add_number(doc, "files 表 / 存储位置确认（MVP 本地磁盘路径 app/storage/{yyyy}/{mm}/）；")
    add_number(doc, "audit_log() 函数签名与 7.2 节一致；")
    add_number(doc, "iter_messages / iter_audit_logs 流式实现确认；")
    add_number(doc, "zipstream-ng 异步流式 zip 实际可用性验证；")

    add_page_break(doc)


def write_appendix(doc):
    add_h1(doc, "附录 A  术语表")
    add_data_table(doc,
        headers=["术语", "定义"],
        rows=[
            ["Agent", "对员工展示的\"智能体\"单元；内部对应 Dify 的一个应用（apps 表里的一行）"],
            ["App", "内部代码 / DB / API 层使用的术语，与 Agent 含义相同"],
            ["Dify", "底层 LLM 应用编排 + RAG 引擎，本项目复用其社区版"],
            ["SSE", "Server-Sent Events，基于 HTTP 长连接的服务器推送协议；本项目用于流式对话"],
            ["JWT", "JSON Web Token，本项目用作 access / refresh token 载体"],
            ["CSRF", "Cross-Site Request Forgery，跨站请求伪造；本项目用 SameSite + Origin 校验防御"],
            ["MIME", "Multipurpose Internet Mail Extensions；本项目用作文件类型白名单匹配"],
            ["Bcrypt", "密码哈希算法；本项目用 passlib[bcrypt]"],
            ["Fernet", "对称加密库（cryptography.fernet）；本项目用于加密 Dify API Key"],
            ["PBKDF2", "Password-Based Key Derivation Function 2；本项目用作密钥派生函数"],
            ["AES-256-GCM", "对称加密算法；本项目用作导出文件加密"],
            ["Penumbra", "@transcend-io/penumbra 库；浏览器侧流式 AES-GCM 加密"],
            ["zxcvbn", "DropBox 开源密码强度评估库；纯前端"],
            ["Envelope", "本项目自定义的 .enc 文件格式：JSON 头 + \\n + 二进制加密体"],
            ["审计三态", "流式端点的审计语义：完成 / 客户端断开 / 服务端异常"],
            ["MCP", "Model Context Protocol（预留）；本期不实现，接口预留"],
        ],
        col_widths=[3.0, 12.3],
        first_col_bold=True,
    )

    add_page_break(doc)

    add_h1(doc, "附录 B  参考文献")
    add_bullet(doc, "企业内部 Agent 平台设计文档（docs/superpowers/specs/2026-08-28-agent-platform-design.md）")
    add_bullet(doc, "MVP Phase 1 实施计划（docs/superpowers/plans/2026-08-28-mvp-phase-1-infra-auth.md）")
    add_bullet(doc, "实施环境准备文档（docs/superpowers/plans/2026-08-28-environment-setup.md）")
    add_bullet(doc, "加密导出特性设计文档（docs/superpowers/specs/2026-08-28-encrypted-export-design.md）")
    add_bullet(doc, "Dify 官方文档：https://docs.dify.ai/")
    add_bullet(doc, "FastAPI 官方文档：https://fastapi.tiangolo.com/")
    add_bullet(doc, "OWASP Password Storage Cheat Sheet（PBKDF2 推荐轮数）")
    add_bullet(doc, "Web Crypto API：https://developer.mozilla.org/en-US/docs/Web/API/Web_Crypto_API")
    add_bullet(doc, "PostgreSQL 16 官方文档：https://www.postgresql.org/docs/16/")
    add_bullet(doc, "SQLAlchemy 2.0 异步文档：https://docs.sqlalchemy.org/en/20/orm/extensions/asyncio.html")

    # 文档结束
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=24, after=0, line_spacing=1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 全文完 ——")
    set_run_font(run, size_pt=12, bold=True, color="888888")


# ============================================================================
# 8. 主入口
# ============================================================================

def main():
    output_dir = Path("/mnt/e/program/agent_platform/output")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "Agent平台建设技术方案框架.docx"

    doc = init_document()

    # 封面
    add_cover(doc)

    # 目录
    add_toc(doc)

    # 章节内容
    write_chapter_1(doc)
    write_chapter_2(doc)
    write_chapter_3(doc)
    write_chapter_4(doc)
    write_chapter_5(doc)
    write_chapter_6(doc)
    write_chapter_7(doc)
    write_chapter_8(doc)
    write_chapter_9(doc)
    write_appendix(doc)

    doc.save(str(output_path))
    print(f"✅ Word 文档已生成：{output_path}")
    print(f"   文件大小：{output_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()